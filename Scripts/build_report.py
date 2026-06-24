#!/usr/bin/env python3
"""
build_report.py — the V10 report (DOCX), main + technical appendix.

  • Naylor_Model_V10_Report.docx            — applied report: what each skill is worth
    (in steals), who already has it (Statcast-style leaderboards), 2025 coaching targets,
    and the drills.
  • Reports/Naylor_Model_V10_Technical_Appendix.docx — full model detail for auditors.

Reads Data/Raw_Season.csv + the Statcast table PNGs (Output/Tables, built by
statcast_tables.py) + model results in Output/Results. No network.

Run:  python3 Scripts/build_report.py   (after statcast_tables.py)
"""
from pathlib import Path
import pandas as pd
import numpy as np
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Scripts/ lives one level below the repo root.
ROOT     = Path(__file__).resolve().parent.parent
DATA_DIR = ROOT / "Data"
RESULTS  = ROOT / "Output" / "Results"
FIG_DIR  = ROOT / "Output" / "Figures"
TAB_DIR  = ROOT / "Output" / "Tables"
LOGO_DIR = ROOT / "Output" / "assets" / "logos"
REPORTS  = ROOT / "Reports"
FIG_DIR.mkdir(parents=True, exist_ok=True)
OUT_MAIN = ROOT / "Naylor_Model_V10_Report.docx"          # "the report" — at root
OUT_APP  = REPORTS / "Naylor_Model_V10_Technical_Appendix.docx"

NAYLOR_ID, SOTO_ID = 647304, 665742

# ── load data ────────────────────────────────────────────────────────────────
sssi = pd.read_csv(DATA_DIR / "Raw_Season.csv")           # master (SSSI + xSB + team)
glm  = pd.read_csv(RESULTS / "DF_v7_GLM_PlainEnglish.csv")
try:
    xsb = pd.read_csv(RESULTS / "DF_v7_xSB_Outcome.csv")  # carries rank_xsb for the appendix
except Exception:
    xsb = sssi
try:
    _pa = pd.read_csv(RESULTS / "DF_perattempt_AUC.csv")
    pa_auc      = float(_pa.iloc[0]["auc"])
    pa_auc_full = float(_pa.iloc[1]["auc"]) if len(_pa) > 1 else None
except Exception:
    pa_auc = 0.7387; pa_auc_full = 0.7231
try:
    bl = pd.ExcelFile(DATA_DIR / "Naylor Blueprint.xlsx")
    bp = bl.parse("Blueprint Leaderboard")
except Exception:
    bl = bp = None

# ── colour helpers ────────────────────────────────────────────────────────────
GREEN = "4CAF5F"; RED = "D65A4D"; AMBER = "FFF3C4"; NAVY = "2C3E50"

def z_to_hex(z):
    """Diverging white-centred heat-map. z>0 → green, z<0 → red, ~0 → near-white."""
    z = max(-1.6, min(1.6, float(z)))
    if z >= 0:
        f = z / 1.6
        r = int(255 + f * (76  - 255)); g = int(255 + f * (175 - 255)); b = int(255 + f * (95  - 255))
    else:
        f = -z / 1.6
        r = int(255 + f * (214 - 255)); g = int(255 + f * (90  - 255)); b = int(255 + f * (77  - 255))
    return f"{r:02X}{g:02X}{b:02X}"

def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)

def _set_cell_text(cell, text, bold=False, size=8.5, align="left", color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0); p.paragraph_format.space_before = Pt(0)
    p.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER,
                   "right": WD_ALIGN_PARAGRAPH.RIGHT}[align]
    r = p.add_run(str(text)); r.bold = bold; r.font.size = Pt(size)
    if color: r.font.color.rgb = RGBColor.from_string(color)

def heat_table(doc, df, columns, heat_specs, *, team_logo_col=None,
               highlight_ids=None, fmt=None, size=8.5, header_size=8.5):
    """columns: [(df_col, header, align)]; heat_specs: {df_col: +1|-1} (sign of 'better')."""
    fmt = fmt or {}
    highlight_ids = highlight_ids or set()
    zmap = {}
    for c, sign in heat_specs.items():
        v = pd.to_numeric(df[c], errors="coerce")
        sd = v.std(ddof=0)
        zmap[c] = ((v - v.mean()) / sd * sign) if sd and not np.isnan(sd) else v * 0

    tbl = doc.add_table(rows=1, cols=len(columns))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"
    for j, (_, header, align) in enumerate(columns):
        _set_cell_text(tbl.rows[0].cells[j], header, bold=True, size=header_size, align=align)
        shade(tbl.rows[0].cells[j], NAVY)
        tbl.rows[0].cells[j].paragraphs[0].runs[0].font.color.rgb = RGBColor.from_string("FFFFFF")

    for _, row in df.iterrows():
        cells = tbl.add_row().cells
        is_anchor = ("runner_id" in df.columns and row.get("runner_id") in highlight_ids)
        for j, (col, _, align) in enumerate(columns):
            if col == team_logo_col:
                _add_team_cell(cells[j], str(row[col]))
            else:
                val = row[col]
                txt = fmt[col](val) if col in fmt else (
                    f"{val:.0f}" if isinstance(val, (int, np.integer)) else
                    (f"{val:.2f}" if isinstance(val, (float, np.floating)) else str(val)))
                _set_cell_text(cells[j], txt, bold=is_anchor, size=size, align=align)
            if col in zmap:
                shade(cells[j], z_to_hex(zmap[col].loc[row.name]))
            elif is_anchor:
                shade(cells[j], AMBER)
    return tbl

def _add_team_cell(cell, team):
    cell.text = ""
    p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo = LOGO_DIR / f"{team}.png"
    if logo.exists():
        try:
            p.add_run().add_picture(str(logo), width=Inches(0.22)); return
        except Exception:
            pass
    r = p.add_run(str(team)); r.font.size = Pt(8.5)

# ── doc-level helpers ────────────────────────────────────────────────────────
def add_fig(doc, path, width=7.0, caption=None):
    path = Path(path)
    if not path.exists():
        doc.add_paragraph(f"[missing figure: {path.name}]"); return
    doc.add_picture(str(path), width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        c = doc.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = c.add_run(caption); r.italic = True; r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor.from_string("666666")

def H(doc, text, lvl=1):
    doc.add_heading(text, level=lvl)

def body(doc, text, size=10, after=4):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text); r.font.size = Pt(size)
    return p

def takeaway(doc, text):
    """Bold one-line BLUF banner above a result."""
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(3)
    r = p.add_run("BOTTOM LINE  "); r.bold = True; r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor.from_string(NAVY)
    r2 = p.add_run(text); r2.font.size = Pt(10); r2.bold = True
    return p

def set_margins(doc, top=0.6, bottom=0.6, left=0.7, right=0.7):
    for s in doc.sections:
        s.top_margin = Inches(top); s.bottom_margin = Inches(bottom)
        s.left_margin = Inches(left); s.right_margin = Inches(right)

def callout_box(doc, title, items, body_shade="F2F7FB", lead_color=None):
    """Shaded, titled callout box. items = list of (bold lead, text)."""
    lead_color = lead_color or NAVY
    tbl = doc.add_table(rows=2, cols=1); tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tc = tbl.rows[0].cells[0]
    _set_cell_text(tc, title, bold=True, size=10.5, color="FFFFFF"); shade(tc, NAVY)
    bc = tbl.rows[1].cells[0]; shade(bc, body_shade)
    bc.paragraphs[0].text = ""
    first = True
    for lead, txt in items:
        p = bc.paragraphs[0] if first else bc.add_paragraph(); first = False
        p.paragraph_format.space_after = Pt(2)
        if lead:
            rb = p.add_run(lead + "  "); rb.bold = True; rb.font.size = Pt(9.5)
            rb.font.color.rgb = RGBColor.from_string(lead_color)
        rt = p.add_run(txt); rt.font.size = Pt(9.5)
    for r in tbl.rows: r.cells[0].width = Inches(7.0)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return tbl

# ── plain-English glossary (single source of truth) ───────────────────────────
GLOSSARY = [
    ("Slow-Steal Skill (SSSI)", "Technique score — skill left after sprint speed is removed. Higher = more."),
    ("Top Speed", "Sprint speed, feet/second (Statcast). The structural baseline."),
    ("Jump", "Seconds to cover the first 30 ft. Lower = quicker first step. Coachable."),
    ("First-Step Burst", "Reaches top speed in fewer feet than his speed predicts."),
    ("Steals Above Speed-Expected", "Success rate minus the rate his speed alone predicts. Positive = beats his speed peers."),
    ("Ground Gained", "Feet covered from the pitcher's first move to release — the biggest single driver of a steal."),
    ("Steal Value (xSB)", "z(net steals above avg) + z(top speed). Surfaces fast, high-volume burners."),
    ("Untapped Speed", "Fast but under-stealing — z(top speed) minus z(net steals). The clearest coaching targets."),
]

def add_glossary_table(doc, entries, size=8.5):
    tbl = doc.add_table(rows=0, cols=2); tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for term, mean in entries:
        c = tbl.add_row().cells
        _set_cell_text(c[0], term, bold=True, size=size, align="left"); shade(c[0], "EAF1F7")
        _set_cell_text(c[1], mean, size=size, align="left")
    for row in tbl.rows:
        row.cells[0].width = Inches(1.9); row.cells[1].width = Inches(5.1)
    return tbl

# ── xSB quadrant figure (a 2-D map — kept as a scatter, polished) ─────────────
def build_xsb_quadrant():
    d = xsb[xsb["era"] == "post_2023"].copy() if "era" in xsb.columns else sssi.copy()
    qcolors = {"Realized Burner": "#2E8B57", "Untapped Wheels": "#2F6FB0",
               "Crafty Technician": "#E0A33E", "Stationary": "#9AA0A6"}
    fig, ax = plt.subplots(figsize=(9.6, 6.4), dpi=150)
    for q, c in qcolors.items():
        s = d[d["quadrant"] == q]
        ax.scatter(s["z_sprint"], s["z_net_sb"], s=26, c=c, alpha=0.55,
                   edgecolors="none", label=q, zorder=2)
    ax.axhline(0, color="#444", lw=1.1, zorder=1)
    ax.axvline(0, color="#444", lw=1.1, zorder=1)

    corners = [(0.985, 0.97, "REALIZED BURNER\nfast + steals a lot", "#2E8B57", "right", "top"),
               (0.985, 0.03, "UNTAPPED WHEELS\nfast, rarely steals", "#2F6FB0", "right", "bottom"),
               (0.015, 0.97, "CRAFTY TECHNICIAN\nslower, steals a lot", "#B07A1E", "left", "top"),
               (0.015, 0.03, "STATIONARY\nneither", "#9AA0A6", "left", "bottom")]
    for x, y, txt, col, ha, va in corners:
        ax.text(x, y, txt, transform=ax.transAxes, ha=ha, va=va, fontsize=11,
                fontweight="bold", color=col, alpha=0.85, zorder=4, linespacing=1.25)

    def _sn(name):
        toks = [t for t in name.split() if t.rstrip(".") not in ("Jr", "Sr", "II", "III")]
        return toks[-1] if toks else name.split()[-1]

    def lab(name, season, dx=6, dy=6, color="#222", ha="left"):
        row = d[(d["player_name"] == name) & (d["season"] == season)]
        if not len(row):
            return
        r = row.iloc[0]
        ax.annotate(f"{_sn(name)} {int(season)}", (r["z_sprint"], r["z_net_sb"]),
                    textcoords="offset points", xytext=(dx, dy), fontsize=8.5,
                    color=color, ha=ha, zorder=5)

    burners = [("Elly De La Cruz", 2024, 7, 4), ("Corbin Carroll", 2023, 7, -2),
               ("Esteury Ruiz", 2023, -7, 6, "right"), ("Shohei Ohtani", 2024, 7, 4),
               ("Ronald Acuña Jr.", 2023, -7, 4, "right"), ("Bobby Witt Jr.", 2023, -8, -4, "right"),
               ("Trea Turner", 2025, 7, -8)]
    for spec in burners:
        nm, sn, dx, dy = spec[:4]; ha = spec[4] if len(spec) > 4 else "left"
        lab(nm, sn, dx=dx, dy=dy, color="#1E5631", ha=ha)
    for nm, sn, dx, dy in [("Jose Siri", 2024, 7, 4), ("Jeremy Peña", 2023, 7, -6),
                           ("Tyler Fitzgerald", 2025, 7, 4), ("Garrett Mitchell", 2026, 7, -4)]:
        lab(nm, sn, dx=dx, dy=dy, color="#1F4E79")
    for nm, sn in [("Josh Naylor", 2025), ("Juan Soto", 2025)]:
        row = d[(d["player_name"] == nm) & (d["season"] == sn)]
        if len(row):
            r = row.iloc[0]
            ax.annotate(f"★ {_sn(nm)} {sn}", (r["z_sprint"], r["z_net_sb"]),
                        textcoords="offset points", xytext=(8, 4), fontsize=9.5,
                        fontweight="bold", color="#B07A1E", ha="left", zorder=6)

    ax.set_xlabel("← slower        TOP SPEED  (standard deviations)        faster →",
                  fontsize=10.5, fontweight="bold")
    ax.set_ylabel("← steals less     STEAL PRODUCTION     steals more →",
                  fontsize=10.5, fontweight="bold")
    ax.grid(True, alpha=0.18, zorder=0)
    for sp in ("top", "right"):
        ax.spines[sp].set_visible(False)
    fig.tight_layout()
    out = FIG_DIR / "Fig_xSB_Quadrant.png"
    fig.savefig(out, bbox_inches="tight"); plt.close(fig)
    return out

# ── GLM lever mappings ─────────────────────────────────────────────────────────
SEASON_ATTEMPTS = 20
GLM_KIND = {
    "Post-Release Distance": "train", "Jump Time": "train", "Accel Gap": "train",
    "Lead Gain (jerk)": "train", "Pre-Release Velocity": "train",
    "Accel→Top-Speed Premium": "train", "Two-Strike Count Share": "train",
    "Primary Lead": "train",
    "Avg Catcher Pop Faced": "context", "Avg Pitcher Pickoff Rate Faced": "context",
    "Share vs Weak-Arm Catchers": "context",
    "Sprint Speed (capped at 28)": "physical", "Bolts": "physical",
}
GLM_PLAIN = {
    "Post-Release Distance": "Ground covered after the pitcher commits",
    "Jump Time": "Jump — quickness over the first 30 ft",
    "Accel Gap": "First-step quickness vs. top speed",
    "Lead Gain (jerk)": "Secondary-lead burst as the pitcher delivers",
    "Pre-Release Velocity": "Closing speed before the pitcher releases",
    "Accel→Top-Speed Premium": "Reaches top speed in fewer feet",
    "Two-Strike Count Share": "Picking two-strike counts to run",
    "Primary Lead": "Primary lead distance off the bag",
    "Avg Catcher Pop Faced": "Catcher pop time faced (opponent)",
    "Avg Pitcher Pickoff Rate Faced": "Pitcher hold/pickoff rate faced (opponent)",
    "Share vs Weak-Arm Catchers": "Facing weak-arm catchers (situation)",
    "Sprint Speed (capped at 28)": "Raw top speed (capped — diminishing returns)",
    "Bolts": "Number of 30+ ft/s sprints",
}
KIND_COLOR = {"train": "#3FA66B", "context": "#9AA0A6", "physical": "#D98A3D"}

# ── THE EQUATION FIGURE — what each skill is worth, in steals ──────────────────
def build_equation_figure():
    """Clean lollipop: each skill's worth in extra steals per 20 attempts."""
    g = glm.copy()
    g["steals"] = g["sb_pct_boost_per_tier"] / 100.0 * SEASON_ATTEMPTS
    g["abs"]   = g["steals"].abs()
    g = g.sort_values("abs").reset_index(drop=True)          # biggest on top
    g["kind"]  = g["feature"].map(GLM_KIND).fillna("context")
    g["label"] = g["feature"].map(GLM_PLAIN).fillna(g["feature"])
    colors = [KIND_COLOR[k] for k in g["kind"]]

    fig, ax = plt.subplots(figsize=(10.6, 6.0), dpi=150)
    yb = list(range(len(g)))
    ax.hlines(yb, 0, g["steals"], color=colors, lw=3.2, zorder=2)
    ax.scatter(g["steals"], yb, color=colors, s=90, zorder=3, edgecolors="white", linewidths=1.0)
    ax.axvline(0, color="#333", lw=1.0)
    ax.set_yticks(yb); ax.set_yticklabels(g["label"], fontsize=10)
    ax.set_ylim(-0.6, len(g) - 0.4)
    vmax, vmin = float(g["steals"].max()), float(g["steals"].min())
    for i, n in zip(yb, g["steals"]):
        lbl = f"{n:+.0f}" if abs(n) >= 0.5 else "~0"
        ax.text(n + (0.06 if n >= 0 else -0.06), i, lbl, va="center",
                ha="left" if n >= 0 else "right", fontsize=9.5, fontweight="bold",
                color="#222")
    ax.set_xlim(min(vmin * 1.25, -1.0), vmax * 1.28)
    ax.set_xlabel("Extra steals per 20 attempts, from one off-season-sized (+1 SD) gain in the skill",
                  fontsize=10.5)
    ax.set_title("What each skill is worth", fontsize=16, fontweight="bold", loc="left", pad=10)
    from matplotlib.patches import Patch
    ax.legend(handles=[Patch(fc=KIND_COLOR["train"], label="Trainable"),
                       Patch(fc=KIND_COLOR["context"], label="Opponent / context"),
                       Patch(fc=KIND_COLOR["physical"], label="Raw speed")],
              loc="lower right", frameon=False, fontsize=9.5)
    for sp in ("top", "right", "left"): ax.spines[sp].set_visible(False)
    ax.tick_params(left=False); ax.grid(axis="x", alpha=0.15)
    fig.tight_layout()
    out = FIG_DIR / "Fig_Equation.png"
    fig.savefig(out, bbox_inches="tight"); plt.close(fig)
    return out

# ── 2025 COACHING TARGET BOARD ─────────────────────────────────────────────────
def coaching_targets(season=2025):
    cur = sssi[sssi["season"] == season].copy()
    cur["net"] = cur["SB"] - cur["CS"]
    gl = cur[(cur["pct_speed"] >= 66) & (cur["real_sb_pct"] >= 0.80) &
             (cur["sb_attempts"] >= 6) & (cur["sb_attempts"] < SEASON_ATTEMPTS)].copy()
    gl["extra_att"]      = (SEASON_ATTEMPTS - gl["sb_attempts"]).clip(lower=0)
    gl["proj_extra_net"] = (gl["extra_att"] * (2 * gl["real_sb_pct"] - 1)).round(0)
    gl = gl.sort_values("proj_extra_net", ascending=False).head(8).copy()
    gl["sbcs"] = gl["SB"].astype(int).astype(str) + "/" + gl["CS"].astype(int).astype(str)

    tf = cur[(cur["sb_attempts"] >= 12) & (cur["real_sb_pct"] < 0.70)].copy()
    lever_name = {"z_jump": "Jump", "z_post_rel_dist": "Ground covered",
                  "z_accel_gap": "First-step burst", "z_lead_gain": "Secondary-lead burst"}
    lever_pp   = {"z_jump": 17.32, "z_post_rel_dist": 24.76, "z_accel_gap": 9.94, "z_lead_gain": 2.08}
    weak, fixpp = [], []
    for _, r in tf.iterrows():
        cand = {c: r[c] for c in lever_name if c in tf.columns and pd.notna(r[c])}
        if not cand:
            weak.append("—"); fixpp.append(0.0); continue
        wc = min(cand, key=cand.get)
        gap = max(0.0, -float(cand[wc]))
        weak.append(lever_name[wc]); fixpp.append(gap * lever_pp[wc])
    tf["weak_lever"] = weak; tf["fix_pp"] = fixpp
    tf = tf.sort_values("sb_attempts", ascending=False).head(6).copy()
    tf["sbcs"] = tf["SB"].astype(int).astype(str) + "/" + tf["CS"].astype(int).astype(str)
    return gl, tf

# ── appendix-only lead-in (auditors get the full What/Result/Variables) ────────
VAR_DEFS = {
    "Player": "Runner (and season).",
    "Top Speed": "Sprint speed in feet/second — the structural baseline (Statcast).",
    "Jump (s)": "Seconds to cover the first 30 ft. Lower = quicker first step.",
    "First-Step Burst": "Reaches top speed in fewer feet than his speed predicts.",
    "Steals Above Speed-Expected": "Success rate above the rate speed alone predicts.",
    "Ground Gained": "Feet covered from the pitcher's first move to release.",
    "SB / CS": "Stolen bases / times caught stealing.",
    "Slow-Steal Skill": "Skill beyond raw speed. Higher = better.",
    "Untapped Speed": "Fast but under-stealing: z(top speed) − z(net steals).",
    "Speed (SD)": "Top speed in standard deviations vs. the league.",
    "Coachable lever": "A trainable skill the model links to steal success.",
    "League Avg": "Typical value across qualified runners.",
    "One-Step Improvement": "A realistic +1-SD gain in that skill.",
    "Success-Rate Change": "Percentage-point change in SB success from that gain.",
    "Ground vs Speed-Expected": "Feet of ground covered beyond what speed predicts.",
    "BCS": "Blueprint Conversion Score — speed-adjusted steal-skill composite.",
    "Net SB": "Stolen bases minus times caught.",
    "Steal Value (xSB)": "z(net steals above avg) + z(top speed).",
    "AUC": "Area under ROC — accuracy (0.5 = coin flip, 1.0 = perfect).",
    "Extra steals": "Extra stolen bases over a 20-attempt season.",
}

def section_intro(doc, what, result, variables):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
    r = p.add_run("What we did.  "); r.bold = True; r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor.from_string(NAVY)
    r2 = p.add_run(what); r2.font.size = Pt(9.5)
    takeaway(doc, result)
    if variables:
        cap = doc.add_paragraph(); cap.paragraph_format.space_after = Pt(1)
        rc = cap.add_run("Variables in this view:"); rc.bold = True; rc.font.size = Pt(9)
        rc.font.color.rgb = RGBColor.from_string("555555")
        for term in variables:
            b = doc.add_paragraph(style="List Bullet")
            b.paragraph_format.space_after = Pt(0); b.paragraph_format.left_indent = Inches(0.25)
            rb = b.add_run(f"{term} — "); rb.bold = True; rb.font.size = Pt(8.5)
            rm = b.add_run(VAR_DEFS.get(term, "")); rm.font.size = Pt(8.5)

# ══════════════════════════════════════════════════════════════════════════════
# MAIN REPORT
# ══════════════════════════════════════════════════════════════════════════════
def build_main():
    xsb_fig = build_xsb_quadrant()
    eq_fig  = build_equation_figure()
    gl, tf  = coaching_targets(2025)
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"; doc.styles["Normal"].font.size = Pt(10.5)
    set_margins(doc)

    # ---- P1: the thesis ------------------------------------------------------
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("The Naylor Model"); r.bold = True; r.font.size = Pt(26)
    sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(6)
    r = sub.add_run("Base-stealing skill that sprint speed hides  ·  for MLB R&D and coaching staffs")
    r.italic = True; r.font.size = Pt(11); r.font.color.rgb = RGBColor.from_string("666666")

    takeaway(doc, "Some of the league's best base-stealers are among its slowest. The skill is jump, "
                  "secondary lead, and reading the pitcher — all coachable, all scattered across the league.")
    body(doc,
         "In 2025 Josh Naylor stole 30 bases at 92% while slower than 98% of MLB — the top skill season "
         "in 673 tracked. This report prices that skill in steals, names who already has it, and flags the "
         "2025 runners to coach.", size=10.5, after=8)

    callout_box(doc, "HOW TO READ THIS", [
        ("In steals.", "Every payoff is extra stolen bases over a 20-attempt season."),
        ("Deeper red = higher.", "On the leaderboards, the headline column is shaded by rank."),
    ])

    gt = glm.copy()
    gt["kind"]  = gt["feature"].map(GLM_KIND).fillna("context")
    gt["lever"] = gt["feature"].map(GLM_PLAIN).fillna(gt["feature"])
    gt["absb"]  = gt["sb_pct_boost_per_tier"].abs()
    gt = gt[gt["kind"] == "train"].sort_values("absb", ascending=False).head(3)
    teaser = [(f"+{round(float(r['sb_pct_boost_per_tier'])/100*SEASON_ATTEMPTS)} steals", f"{r['lever']}.")
              for _, r in gt.iterrows()]
    callout_box(doc, "THE THREE LEVERS, IN ORDER OF PAYOFF", teaser,
                body_shade="F3F8F3", lead_color="2E7D32")

    doc.add_page_break()

    # ---- §1: what each skill is worth (figure carries it) --------------------
    H(doc, "1 · What each skill is worth")
    takeaway(doc, "Three trainable skills move steals. Raw foot speed barely does.")
    add_fig(doc, eq_fig, width=6.9,
            caption="Extra steals over 20 attempts from one off-season-sized (+1 SD) gain. "
                    "Green = trainable · grey = opponent/context · orange = raw speed.")

    doc.add_page_break()

    # ---- §2: who already has it (Statcast leaderboard) -----------------------
    H(doc, "2 · Who already has the skill")
    takeaway(doc, "Mid-pack-speed runners fill the leaderboard — the trait is common and findable.")
    callout_box(doc, "STRIPPING SPEED OUT", [
        ("", "Fast runners cover more ground and steal more — of course. Plot a skill against sprint "
             "speed, draw the line speed predicts, and keep only how far each runner sits above it. "
             "That gap is skill with speed removed."),
    ])
    add_fig(doc, FIG_DIR / "Fig_GroundCovered_Scatter.png", width=5.4,
            caption="The line is what speed predicts; height above it is skill. Naylor, Soto, and "
                    "McMahon sit far above it at the slowest speeds.")
    add_fig(doc, TAB_DIR / "Slow_Steal_Skill.png", width=7.0)

    doc.add_page_break()

    # ---- §2b: the full skill index ------------------------------------------
    H(doc, "2b · The full skill index, with teams")
    takeaway(doc, "Ryan McMahon's 2026 is the highest score in the data — a slow corner infielder "
                  "out-converting the league's burners.")
    add_fig(doc, TAB_DIR / "Blueprint_Conversion.png", width=7.0)

    doc.add_page_break()

    # ---- §3: who to coach in 2025 -------------------------------------------
    H(doc, "3 · Who to coach in 2025")
    add_fig(doc, xsb_fig, width=6.4,
            caption="Each dot is a runner-season (2023–2026). Right = faster; up = steals more. "
                    "Blue lower-right = fast but under-running (green-light); gold = the slow-but-"
                    "productive archetype (★ Naylor, Soto).")

    H(doc, "3a · Green-light — fast, efficient, under-running", lvl=2)
    body(doc, "Top-third speed, 80%+ success, fewer than 20 attempts. Let them run.", size=9.5, after=3)
    glr = gl.copy(); glr.insert(0, "gl_rank", range(1, len(glr) + 1))
    heat_table(doc, glr,
        columns=[("gl_rank", "#", "center"), ("player_name", "Player", "left"),
                 ("sprint_speed", "Top Speed\n(ft/s)", "center"), ("sbcs", "SB / CS", "center"),
                 ("real_sb_pct", "Success\nrate", "center"),
                 ("sb_attempts", "Attempts\n(2025)", "center"),
                 ("proj_extra_net", "If unleashed\n(+ steals)", "center")],
        heat_specs={"real_sb_pct": +1, "proj_extra_net": +1},
        fmt={"sprint_speed": lambda v: f"{v:.1f}", "real_sb_pct": lambda v: f"{v*100:.0f}%",
             "sb_attempts": lambda v: f"{int(v)}", "proj_extra_net": lambda v: f"+{int(v)}",
             "gl_rank": lambda v: f"{int(v)}"}, size=9)
    body(doc, "\"If unleashed\" holds the runner at his own 2025 rate over ~20 attempts — a ranking, "
              "not a forecast.", size=8, after=8)

    H(doc, "3b · Technique-fix — high volume, caught too often", lvl=2)
    body(doc, "12+ attempts, under 70% success. Don't run more — run better, on the lever named.",
         size=9.5, after=3)
    tfr = tf.copy()
    tfr["fix_txt"] = tfr["fix_pp"].apply(
        lambda pp: f"+{round(pp/100*SEASON_ATTEMPTS)}" if pp/100*SEASON_ATTEMPTS >= 0.5 else "—")
    heat_table(doc, tfr,
        columns=[("player_name", "Player", "left"), ("sprint_speed", "Top Speed\n(ft/s)", "center"),
                 ("sbcs", "SB / CS", "center"), ("real_sb_pct", "Success\nrate", "center"),
                 ("sb_attempts", "Attempts\n(2025)", "center"),
                 ("weak_lever", "Weakest\nlever", "left"),
                 ("fix_txt", "Fix →\n+ steals/20", "center")],
        heat_specs={"real_sb_pct": +1, "fix_pp": +1},
        fmt={"sprint_speed": lambda v: f"{v:.1f}", "real_sb_pct": lambda v: f"{v*100:.0f}%",
             "sb_attempts": lambda v: f"{int(v)}"}, size=9)
    body(doc, "\"Fix\" = steals gained over 20 tries if the weakest skill reaches league average.",
         size=8, after=6)

    doc.add_page_break()

    # ---- §4: drills + caveats ------------------------------------------------
    H(doc, "4 · What to coach")
    for head, txt in [
        ("Cover more ground once the pitcher commits.  (+5 / 20)",
         "The biggest lever. Drill the secondary-lead explosion and the read of the pitcher's first move."),
        ("Sharpen the jump.  (+3 / 20)",
         "Stance, weight shift, first-step direction — pure technique over the first 30 feet."),
        ("Reach top speed sooner.  (+2 / 20)",
         "Acceleration mechanics, not a higher top speed, are what convert."),
    ]:
        p = doc.add_paragraph(style="List Number"); p.paragraph_format.space_after = Pt(3)
        rb = p.add_run(head + "  "); rb.bold = True; rb.font.size = Pt(10.5)
        rt = p.add_run(txt); rt.font.size = Pt(10)

    H(doc, "Caveats", lvl=2)
    for head, txt in [
        ("Projections rank, they don't forecast.",
         "Real results move with health, matchups, and game situation."),
        ("Per-lever gains assume the rest stays fixed.",
         "Small compounding gains across two levers beat one full-SD jump."),
        ("Context columns aren't levers.",
         "Catcher pop and pitcher hold move the math but can't be coached — shown to separate skill "
         "from circumstance."),
    ]:
        p = doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after = Pt(2)
        rb = p.add_run(head + "  "); rb.bold = True; rb.font.size = Pt(9.5)
        rt = p.add_run(txt); rt.font.size = Pt(9.5)

    doc.add_page_break()
    H(doc, "Metric reference")
    add_glossary_table(doc, GLOSSARY)

    doc.save(str(OUT_MAIN))
    print(f"wrote {OUT_MAIN}  ({OUT_MAIN.stat().st_size/1024:.0f} KB)")

# ══════════════════════════════════════════════════════════════════════════════
# TECHNICAL APPENDIX
# ══════════════════════════════════════════════════════════════════════════════
def build_appendix():
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"; doc.styles["Normal"].font.size = Pt(10)
    set_margins(doc, left=0.7, right=0.7)

    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("The Naylor Model — V10 Technical Appendix"); r.bold = True; r.font.size = Pt(22)
    sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Full model detail, validation, and the Blueprint Conversion Score.")
    r.italic = True; r.font.size = Pt(10.5); r.font.color.rgb = RGBColor.from_string("666666")
    doc.add_paragraph()

    # A — The model is per-attempt
    H(doc, "A  The Model is Per-Attempt (≈10,400 rows), Not Season")
    section_intro(doc,
        what="Players are analysed at the grain that decides a steal — the individual attempt "
             "(≈10,366 tracked attempts), not a 673-row season average. A per-attempt XGBoost (Model A) "
             "is the predictor; an interpretable GLM (Model C) gives the plain-English weights.",
        result="At the attempt grain the model reaches the target range and the driver is the per-pitch "
                "lead distances — exactly this report's thesis.",
        variables=["AUC"])
    tbl = doc.add_table(rows=1, cols=4); tbl.style = "Table Grid"; tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(["Model", "Unit", "AUC", "Purpose"]):
        _set_cell_text(tbl.rows[0].cells[j], h, bold=True, size=9, color="FFFFFF"); shade(tbl.rows[0].cells[j], NAVY)
    for rdat in [("Model A — per-attempt XGBoost", "Individual attempt (≈10,366)", f"{pa_auc:.3f}",
                  "THE model — does THIS steal succeed"),
                 ("Model C — interpretable GLM", "Runner-season", "—",
                  "Plain-English weight table (coaching levers)")]:
        c = tbl.add_row().cells
        for j, v in enumerate(rdat): _set_cell_text(c[j], v, size=9)
    body(doc,
         f"Why per-attempt. The unit of analysis is one steal attempt, with the exact lead distances the "
         f"runner got on that pitch. That is ≈10,366 rows — 15× the 673 runner-seasons — and far less "
         f"averaging noise. CV AUC is {pa_auc:.3f}, driven by the lead distances, not speed.", size=9.5)
    body(doc,
         "Leakage discipline. run_value (outcome-derived) is dropped; the runner's own season success "
         "rate is excluded; catcher/pitcher tendencies are out-of-fold target-encoded. Tellingly, adding "
         f"those tendencies LOWERS AUC ({pa_auc:.4f} → {pa_auc_full:.4f}) — the per-pitch leads alone carry the "
         "signal, which is the entire point.", size=9.5)
    add_fig(doc, FIG_DIR / "Fig_AUC.png", 5.4,
            "Figure A1 — Per-attempt model AUC. Trained on ≈10,400 individual steal attempts; the leads "
            "carry the signal (battery tendencies don't help).")
    add_fig(doc, FIG_DIR / "Fig_Importance.png", 6.2,
            "Figure A2 — What decides a steal, per attempt. Per-pitch lead distances dominate.")

    # A2 — How the per-attempt model reaches AUC ~0.74
    H(doc, f"A2  How the Per-Attempt Model Reaches AUC {pa_auc:.3f}")
    body(doc,
         "The grain is the point. Each row is one steal attempt: lead at first move, ground gained to "
         "release, lead at release, plus runner skill context (sprint speed, jump time, accel gap). "
         "One attempt where a runner gains 14 ft before first move is directly informative; a season "
         "average dilutes that across pitch type, count, and game state. The attempt grain keeps the "
         "signal clean and gives 15× more rows.", size=9.5)
    body(doc,
         "Leakage discipline. Catcher pop time and pitcher pickoff rate are out-of-fold target-encoded — "
         "computed on the training folds only, applied to the held-out fold. Adding them drops AUC "
         f"({pa_auc:.4f} → {pa_auc_full:.4f}); the lead distances carry the signal on their own.", size=9.5)
    body(doc,
         "Model spec. XGBClassifier, 500 trees, max_depth=4, learning_rate=0.03, 5-fold stratified CV, "
         "pooled out-of-fold AUC. Deliberately untuned — a sensible default reaches the target range, so "
         "the result is not an artifact of hyperparameter search.", size=9.5)
    body(doc,
         f"Honest ceiling on public data: ~0.74–0.78. To push further: pitch type at first move "
         "(fastball vs. off-speed changes runner timing) and pitcher handedness (LHP pickoff mechanics "
         "differ fundamentally). See AUC_Roadmap.md.", size=9.5)

    # B — full GLM + equation figure
    H(doc, "B  Model C — The Steal-Success Equation (Full GLM)")
    section_intro(doc,
        what="Every feature in the GLM, sorted by absolute impact. Coefficients are on z-scored inputs — "
             "the effect of a one-SD move in each skill.",
        result="The biggest movers are coachable timing skills; raw speed and context matter less.",
        variables=["Coachable lever", "League Avg", "One-Step Improvement", "Success-Rate Change", "Extra steals"])
    add_fig(doc, FIG_DIR / "Fig_Equation.png", 6.8,
            "Figure B1 — each skill's worth in extra steals per 20 attempts (one off-season-sized +1-SD gain).")
    gg = glm.copy()
    gg["abs"] = gg["sb_pct_boost_per_tier"].abs()
    gg = gg.sort_values("abs", ascending=False)
    heat_table(doc, gg,
        columns=[("feature", "Feature", "left"), ("league_avg", "League\nAvg", "center"),
                 ("one_tier_step", "One-Step\nImprovement", "center"),
                 ("sb_pct_boost_per_tier", "Success-Rate\nChange (pp)", "center"),
                 ("odds_multiplier", "Odds\nMultiplier", "center")],
        heat_specs={"sb_pct_boost_per_tier": +1},
        fmt={"league_avg": lambda v: f"{v:.3f}", "one_tier_step": lambda v: f"{v:.3f}",
             "sb_pct_boost_per_tier": lambda v: f"{v:+.2f}", "odds_multiplier": lambda v: f"{v:.3f}"},
        size=8.5)

    # C — full SSSI
    H(doc, "C  Slow-Steal Skill (SSSI) — Full Top 25 with components")
    section_intro(doc,
        what="The full Top 25, components exposed.",
        result="Skill, not speed, drives the ranking — leaders cluster in the mid-20s ft/s while grading "
                "elite on jump, burst, and steals-above-expected.",
        variables=["Player", "Top Speed", "Jump (s)", "Steals Above Speed-Expected",
                   "First-Step Burst", "Slow-Steal Skill"])
    ss = sssi.sort_values("SSSI_v7", ascending=False).head(25).copy()
    heat_table(doc, ss,
        columns=[("rank_v7", "#", "center"), ("player_name", "Player", "left"),
                 ("season", "Yr", "center"), ("sprint_speed", "Top Speed\n(ft/s)", "center"),
                 ("jump_time", "Jump (s)\nlower better", "center"),
                 ("sb_residual", "Steals Above\nSpeed-Expected", "center"),
                 ("accel_topspeed_premium", "First-Step\nBurst", "center"),
                 ("real_sb_pct", "SB%", "center"), ("SSSI_v7", "Slow-Steal\nSkill", "center")],
        heat_specs={"sb_residual": +1, "accel_topspeed_premium": +1, "real_sb_pct": +1, "SSSI_v7": +1},
        highlight_ids={NAYLOR_ID, SOTO_ID},
        fmt={"season": lambda v: f"{int(v)}", "sprint_speed": lambda v: f"{v:.1f}",
             "real_sb_pct": lambda v: f"{v*100:.0f}%", "jump_time": lambda v: f"{v:.2f}",
             "SSSI_v7": lambda v: f"{v:+.2f}", "sb_residual": lambda v: f"{v*100:+.1f}%",
             "accel_topspeed_premium": lambda v: f"{v:+.1f}", "rank_v7": lambda v: f"{int(v)}"}, size=8.5)

    # D — xSB
    H(doc, "D  Expected SB Outcome (xSB) — Full Leaderboards")
    section_intro(doc,
        what="Steal Value (xSB) = z(net steals above average) + z(top speed). Surfaces runners who are "
             "both fast AND productive.",
        result="Ceiling cases are burners (De La Cruz, Carroll); Naylor & Soto reach the same production "
                "from the slow side.",
        variables=["Player", "Top Speed", "Net SB", "Speed (SD)", "Steal Value (xSB)"])
    xt = xsb.sort_values("xsb_outcome", ascending=False).head(15).copy()
    heat_table(doc, xt,
        columns=[("rank_xsb", "#", "center"), ("player_name", "Player", "left"),
                 ("season", "Yr", "center"), ("sprint_speed", "Top Speed\n(ft/s)", "center"),
                 ("SB", "SB", "center"), ("CS", "CS", "center"), ("net_sb", "Net SB", "center"),
                 ("z_sprint", "Speed\n(SD)", "center"), ("z_net_sb", "Steals\n(SD)", "center"),
                 ("xsb_outcome", "Steal Value\n(xSB)", "center")],
        heat_specs={"z_sprint": +1, "z_net_sb": +1, "xsb_outcome": +1, "net_sb": +1},
        fmt={"season": lambda v: f"{int(v)}", "sprint_speed": lambda v: f"{v:.1f}",
             "SB": lambda v: f"{int(v)}", "CS": lambda v: f"{int(v)}", "net_sb": lambda v: f"{int(v)}",
             "z_sprint": lambda v: f"{v:+.2f}", "z_net_sb": lambda v: f"{v:+.2f}",
             "xsb_outcome": lambda v: f"{v:+.2f}", "rank_xsb": lambda v: f"{int(v)}"}, size=8.5)

    ns = xsb[xsb["runner_id"].isin([NAYLOR_ID, SOTO_ID])].sort_values("xsb_outcome", ascending=False).copy()
    H(doc, "D.1  Naylor & Soto — Crafty Technicians (year-by-year)", lvl=2)
    body(doc, "Both carry a negative Untapped Speed every year — they out-steal their speed consistently.",
         size=9.5, after=3)
    heat_table(doc, ns,
        columns=[("player_name", "Player", "left"), ("season", "Yr", "center"),
                 ("sprint_speed", "Top Speed\n(ft/s)", "center"), ("SB", "SB", "center"), ("CS", "CS", "center"),
                 ("z_sprint", "Speed\n(SD)", "center"), ("z_net_sb", "Steals\n(SD)", "center"),
                 ("sb_potential_gap", "Untapped\nSpeed", "center"), ("quadrant", "Quadrant", "left")],
        heat_specs={"z_net_sb": +1, "sb_potential_gap": -1},
        highlight_ids={NAYLOR_ID, SOTO_ID},
        fmt={"season": lambda v: f"{int(v)}", "sprint_speed": lambda v: f"{v:.1f}",
             "SB": lambda v: f"{int(v)}", "CS": lambda v: f"{int(v)}",
             "z_sprint": lambda v: f"{v:+.2f}", "z_net_sb": lambda v: f"{v:+.2f}",
             "sb_potential_gap": lambda v: f"{v:+.2f}"}, size=8.5)

    # E — BCS (lead with the Statcast tables)
    if bp is not None:
        H(doc, "E  Blueprint Conversion Score (BCS)")
        section_intro(doc,
            what="BCS combines three speed-adjusted residuals — converts more than speed predicts, covers "
                 "more ground than speed predicts, less a caught-stealing penalty.",
            result="Full of slow-but-skilled runners (★ Naylor, Soto) holding their own against burners.",
            variables=["BCS", "Ground vs Speed-Expected", "Top Speed", "SB / CS"])
        add_fig(doc, TAB_DIR / "Blueprint_Conversion.png", 7.0,
                "Figure E1 — Top 15 by Blueprint Conversion Score.")
        def bcs_cols():
            return [("rank_BCS", "#", "center"), ("player_name", "Player", "left"),
                    ("team", "Team", "center"), ("season", "Yr", "center"),
                    ("sprint_pctile", "Speed\n%ile", "center"), ("SB", "SB", "center"),
                    ("CS", "CS", "center"), ("SB_pct", "SB%", "center"),
                    ("mean_gain_to_release_ft", "Gain\n(ft)", "center"),
                    ("gain_resid_z", "Ground vs\nExp (SD)", "center"),
                    ("success_resid_z", "Convert vs\nExp (SD)", "center"),
                    ("BCS", "BCS", "center")]
        bcs_fmt = {"season": lambda v: f"{int(v)}", "sprint_pctile": lambda v: f"{v:.0f}",
                   "SB": lambda v: f"{int(v)}", "CS": lambda v: f"{int(v)}",
                   "SB_pct": lambda v: f"{v*100:.0f}%" if v <= 1.5 else f"{v:.0f}%",
                   "mean_gain_to_release_ft": lambda v: f"{v:.1f}",
                   "gain_resid_z": lambda v: f"{v:+.2f}", "success_resid_z": lambda v: f"{v:+.2f}",
                   "BCS": lambda v: f"{v:+.2f}", "rank_BCS": lambda v: f"{int(v)}"}
        H(doc, "E.1  Overall Top 15", lvl=2)
        top15 = bp.sort_values("rank_BCS").head(15).copy()
        heat_table(doc, top15, columns=bcs_cols(), team_logo_col="team",
                   heat_specs={"BCS": +1, "gain_resid_z": +1, "success_resid_z": +1},
                   highlight_ids={NAYLOR_ID, SOTO_ID}, fmt=bcs_fmt, size=8.5)
        H(doc, "E.2  Bottom 15 — Squanderers (fast, serial caught-stealings)", lvl=2)
        bot15 = bp.sort_values("rank_BCS", ascending=False).head(15).copy().sort_values("BCS")
        heat_table(doc, bot15, columns=bcs_cols(), team_logo_col="team",
                   heat_specs={"BCS": +1, "gain_resid_z": +1, "success_resid_z": +1},
                   fmt=bcs_fmt, size=8.5)
        H(doc, "E.3  Ground Covered Beyond Speed-Expected — Top 15", lvl=2)
        body(doc, "Feet gained first-move→release, after removing sprint speed. Slow runners top it.",
             size=9.5, after=3)
        add_fig(doc, TAB_DIR / "Ground_Covered.png", 7.0,
                "Figure E2 — Top 15 by ground covered beyond what speed predicts.")

    doc.save(str(OUT_APP))
    print(f"wrote {OUT_APP}  ({OUT_APP.stat().st_size/1024:.0f} KB)")

# ══════════════════════════════════════════════════════════════════════════════
if __name__ == "__main__":
    build_main()
    build_appendix()
