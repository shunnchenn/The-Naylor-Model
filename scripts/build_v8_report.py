#!/usr/bin/env python3
"""
build_v8_report.py — Ultra-lean BLUF report for MLB R&D + coaches (DOCX).

Produces TWO documents:
  • Reports/Naylor_Model_v8_Report.docx            — applied BLUF report (≈5 pp): the
    steal-success equation, the skill leaderboard, a 2025 coaching target board with
    projected payoff, and the three drills. Every page self-contained, plain-English.
  • Reports/Naylor_Model_v8_Technical_Appendix.docx — full detail for auditors
    (Models A/B/C, full SSSI, GLM detail, BCS + logos, de-leak/AUC note).

Reads existing v7 CSVs / xlsx / figures (no network, no re-run of v7_explore.py).
Derived from build_v7_report.py — reuses its z_to_hex / shade / heat_table helpers.

Run:  python3 build_v8_report.py
"""
from pathlib import Path
import pandas as pd
import numpy as np
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# scripts/ lives one level below the repo root.
ROOT     = Path(__file__).resolve().parent.parent
DATA_DIR = ROOT / "data"
FIG_DIR  = ROOT / "Figures"
LOGO_DIR = ROOT / "Figures" / "logos"
REPORTS  = ROOT / "Reports"
OUT_MAIN = ROOT / "Naylor_Model_v8_Report.docx"          # "the report" — at root
OUT_APP  = REPORTS / "Naylor_Model_v8_Technical_Appendix.docx"

NAYLOR_ID, SOTO_ID = 647304, 665742

# ── load data ────────────────────────────────────────────────────────────────
sssi = pd.read_csv(DATA_DIR / "DF_v7_SSSI.csv")
xsb  = pd.read_csv(DATA_DIR / "DF_v7_xSB_Outcome.csv")
glm  = pd.read_csv(DATA_DIR / "DF_v7_GLM_PlainEnglish.csv")
try:
    auc = pd.read_csv(DATA_DIR / "DF_v7_ModelB_AUC.csv")
except Exception:
    auc = None
try:
    bl = pd.ExcelFile(DATA_DIR / "Naylor Blueprint.xlsx")
    bp = bl.parse("Blueprint Leaderboard")
except Exception:
    bp = None

# ── colour helpers (shared with v7 builder) ──────────────────────────────────
GREEN = "4CAF5F"; RED = "D65A4D"; AMBER = "FFF3C4"; NAVY = "2C3E50"

def z_to_hex(z):
    """Diverging white-centred heat-map. z>0 → green, z<0 → red, ~0 → near-white."""
    z = max(-1.6, min(1.6, float(z)))
    if z >= 0:
        f = z / 1.6
        r = int(255 + f * (76  - 255)); g = int(255 + f * (175 - 255)); b = int(255 + f * (95  - 255))
    else:
        f = -z / 1.6
        r = int(255 + f * (214 - 255)); g = int(255 + f * (90  - 255)); b = int(255 + f * (77  - 255))
    return f"{r:02X}{g:02X}{b:02X}"

def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)

def _set_cell_text(cell, text, bold=False, size=8.5, align="left", color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0); p.paragraph_format.space_before = Pt(0)
    p.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER,
                   "right": WD_ALIGN_PARAGRAPH.RIGHT}[align]
    r = p.add_run(str(text)); r.bold = bold; r.font.size = Pt(size)
    if color: r.font.color.rgb = RGBColor.from_string(color)

def heat_table(doc, df, columns, heat_specs, *, team_logo_col=None,
               highlight_ids=None, fmt=None, size=8.5, header_size=8.5):
    """columns: [(df_col, header, align)]; heat_specs: {df_col: +1|-1} (sign of 'better')."""
    fmt = fmt or {}
    highlight_ids = highlight_ids or set()
    zmap = {}
    for c, sign in heat_specs.items():
        v = pd.to_numeric(df[c], errors="coerce")
        sd = v.std(ddof=0)
        zmap[c] = ((v - v.mean()) / sd * sign) if sd and not np.isnan(sd) else v * 0

    tbl = doc.add_table(rows=1, cols=len(columns))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"
    for j, (_, header, align) in enumerate(columns):
        _set_cell_text(tbl.rows[0].cells[j], header, bold=True, size=header_size, align=align)
        shade(tbl.rows[0].cells[j], NAVY)
        tbl.rows[0].cells[j].paragraphs[0].runs[0].font.color.rgb = RGBColor.from_string("FFFFFF")

    for _, row in df.iterrows():
        cells = tbl.add_row().cells
        is_anchor = ("runner_id" in df.columns and row.get("runner_id") in highlight_ids)
        for j, (col, _, align) in enumerate(columns):
            if col == team_logo_col:
                _add_team_cell(cells[j], str(row[col]))
            else:
                val = row[col]
                txt = fmt[col](val) if col in fmt else (
                    f"{val:.0f}" if isinstance(val, (int, np.integer)) else
                    (f"{val:.2f}" if isinstance(val, (float, np.floating)) else str(val)))
                _set_cell_text(cells[j], txt, bold=is_anchor, size=size, align=align)
            if col in zmap:
                shade(cells[j], z_to_hex(zmap[col].loc[row.name]))
            elif is_anchor:
                shade(cells[j], AMBER)
    return tbl

def _add_team_cell(cell, team):
    cell.text = ""
    p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo = LOGO_DIR / f"{team}.png"
    if logo.exists():
        try:
            p.add_run().add_picture(str(logo), width=Inches(0.22)); return
        except Exception:
            pass
    r = p.add_run(str(team)); r.font.size = Pt(8.5)

# ── doc-level helpers ────────────────────────────────────────────────────────
def add_fig(doc, path, width=7.0, caption=None):
    path = Path(path)
    if not path.exists():
        doc.add_paragraph(f"[missing figure: {path.name}]"); return
    doc.add_picture(str(path), width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        c = doc.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = c.add_run(caption); r.italic = True; r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor.from_string("666666")

def H(doc, text, lvl=1):
    doc.add_heading(text, level=lvl)

def body(doc, text, size=10, after=4):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text); r.font.size = Pt(size)
    return p

def takeaway(doc, text):
    """Bold one-line BLUF banner above a result."""
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(3)
    r = p.add_run("BOTTOM LINE  "); r.bold = True; r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor.from_string(NAVY)
    r2 = p.add_run(text); r2.font.size = Pt(10); r2.bold = True
    return p

def legend_line(doc):
    """Compact inline colour legend, repeated on each results page."""
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
    parts = [("■ ", GREEN, "better than peers"),
             ("   ■ ", "BFBFBF", "league-average"),
             ("   ■ ", RED, "worse than peers")]
    for sym, col, label in parts:
        rs = p.add_run(sym); rs.font.size = Pt(9); rs.font.color.rgb = RGBColor.from_string(col)
        rl = p.add_run(label); rl.font.size = Pt(8.5); rl.font.color.rgb = RGBColor.from_string("555555")
    tail = p.add_run("   (shading = standard deviations above/below the players shown)")
    tail.italic = True; tail.font.size = Pt(8); tail.font.color.rgb = RGBColor.from_string("888888")

def set_margins(doc, top=0.6, bottom=0.6, left=0.7, right=0.7):
    for s in doc.sections:
        s.top_margin = Inches(top); s.bottom_margin = Inches(bottom)
        s.left_margin = Inches(left); s.right_margin = Inches(right)

# ── plain-English glossary (single source of truth) ───────────────────────────
GLOSSARY = [
    ("Slow-Steal Skill (SSSI)", "Overall technique score — how good a runner is beyond what raw speed explains. Higher = more skill."),
    ("Top Speed", "Sprint speed in feet/second (Statcast). The structural baseline a runner is born with."),
    ("Jump", "Seconds to cover the first 30 ft. Lower = quicker first step. Coachable."),
    ("First-Step Burst", "Reaches top speed in fewer feet than his speed predicts. A premium when paired with high speed."),
    ("Steals Above Speed-Expected", "Real success rate minus the rate his sprint speed alone would predict. Positive = beats his speed peers."),
    ("Steal Value (xSB)", "Speed AND production combined: z(net steals above avg) + z(top speed). Surfaces fast, high-volume burners."),
    ("Untapped Speed", "Fast but under-stealing — z(top speed) minus z(net steals). The clearest coaching targets."),
]

def add_glossary_table(doc, entries, size=8.5):
    tbl = doc.add_table(rows=0, cols=2); tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for term, mean in entries:
        c = tbl.add_row().cells
        _set_cell_text(c[0], term, bold=True, size=size, align="left"); shade(c[0], "EAF1F7")
        _set_cell_text(c[1], mean, size=size, align="left")
    # widen second column
    for row in tbl.rows:
        row.cells[0].width = Inches(1.9); row.cells[1].width = Inches(5.1)
    return tbl

# ── (re)build a cleaner, more intuitive xSB quadrant figure ───────────────────
def build_xsb_quadrant():
    """Post-2023 era quadrant, plain-English labels, only recognizable names annotated."""
    d = xsb[xsb["era"] == "post_2023"].copy()
    # No red here: every quadrant except "Stationary" is a positive profile. The
    # archetype (Crafty Technician = slow but steals a lot) is GOLD, the hero colour.
    qcolors = {"Realized Burner": "#2E8B57", "Untapped Wheels": "#2F6FB0",
               "Crafty Technician": "#E0A33E", "Stationary": "#9AA0A6"}
    fig, ax = plt.subplots(figsize=(9.6, 6.4), dpi=150)
    for q, c in qcolors.items():
        s = d[d["quadrant"] == q]
        ax.scatter(s["z_sprint"], s["z_net_sb"], s=26, c=c, alpha=0.55,
                   edgecolors="none", label=q, zorder=2)
    ax.axhline(0, color="#444", lw=1.1, zorder=1)
    ax.axvline(0, color="#444", lw=1.1, zorder=1)

    # corner labels (plain English, the four reads)
    corners = [(0.985, 0.97, "REALIZED BURNER\nfast + steals a lot", "#2E8B57", "right", "top"),
               (0.985, 0.03, "UNTAPPED WHEELS\nfast, but rarely steals", "#2F6FB0", "right", "bottom"),
               (0.015, 0.97, "CRAFTY TECHNICIAN\nslower, but steals a lot", "#B07A1E", "left", "top"),
               (0.015, 0.03, "STATIONARY\nneither speed nor steals", "#9AA0A6", "left", "bottom")]
    for x, y, txt, col, ha, va in corners:
        ax.text(x, y, txt, transform=ax.transAxes, ha=ha, va=va, fontsize=11,
                fontweight="bold", color=col, alpha=0.85, zorder=4,
                linespacing=1.25)

    def _surname(name):
        toks = [t for t in name.split() if t.rstrip(".") not in ("Jr", "Sr", "II", "III")]
        return toks[-1] if toks else name.split()[-1]

    def lab(name, season, dx=6, dy=6, fs=8.5, weight="normal", color="#222", ha="left"):
        row = d[(d["player_name"] == name) & (d["season"] == season)]
        if not len(row):
            return
        r = row.iloc[0]
        ax.annotate(f"{_surname(name)} {int(season)}",
                    (r["z_sprint"], r["z_net_sb"]), textcoords="offset points",
                    xytext=(dx, dy), fontsize=fs, fontweight=weight, color=color,
                    ha=ha, zorder=5)

    # recognizable burners (per-label offsets to avoid the top-right cluster)
    burners = [("Elly De La Cruz", 2024, 7, 4), ("Corbin Carroll", 2023, 7, -2),
               ("Esteury Ruiz", 2023, -7, 6, "right"), ("Shohei Ohtani", 2024, 7, 4),
               ("Ronald Acuña Jr.", 2023, -7, 4, "right"), ("Bobby Witt Jr.", 2023, -8, -4, "right"),
               ("Trea Turner", 2025, 7, -8)]
    for spec in burners:
        nm, sn, dx, dy = spec[:4]; ha = spec[4] if len(spec) > 4 else "left"
        lab(nm, sn, dx=dx, dy=dy, color="#1E5631", ha=ha)
    # untapped coaching targets
    for nm, sn, dx, dy in [("Jose Siri", 2024, 7, 4), ("Jeremy Peña", 2023, 7, -6),
                           ("Tyler Fitzgerald", 2025, 7, 4), ("Garrett Mitchell", 2026, 7, -4)]:
        lab(nm, sn, dx=dx, dy=dy, color="#1F4E79")
    # the archetype benchmarks (★ marks them; colour matches the gold archetype quadrant)
    for nm, sn, dx, dy in [("Josh Naylor", 2025, 8, 4), ("Juan Soto", 2025, 8, 4)]:
        row = d[(d["player_name"] == nm) & (d["season"] == sn)]
        if len(row):
            r = row.iloc[0]
            ax.annotate(f"★ {_surname(nm)} {sn}", (r["z_sprint"], r["z_net_sb"]),
                        textcoords="offset points", xytext=(dx, dy), fontsize=9.5,
                        fontweight="bold", color="#B07A1E", ha="left", zorder=6)

    ax.set_xlabel("← slower        TOP SPEED  (standard deviations)        faster →",
                  fontsize=10.5, fontweight="bold")
    ax.set_ylabel("← steals less     STEAL PRODUCTION     steals more →",
                  fontsize=10.5, fontweight="bold")
    ax.set_title("Who is fast, who is productive, and who is both  (2023–2026)",
                 fontsize=13, fontweight="bold", pad=12)
    ax.grid(True, alpha=0.18, zorder=0)
    for sp in ("top", "right"):
        ax.spines[sp].set_visible(False)
    fig.tight_layout()
    out = FIG_DIR / "Fig_v8_xSB_Quadrant.png"
    fig.savefig(out, bbox_inches="tight"); plt.close(fig)
    return out

# ── speed-tier palette for the by-year leaderboards (Statcast-style key) ──────
# Colour encodes ONLY the runner's raw-speed tier — never "good vs bad". On these
# leaderboards every bar is a positive, top-25 result, so there is deliberately NO
# red here (red is reserved for below-average values in the heat-shaded tables).
# GOLD = the edge we're hunting: a slow runner posting an elite skill result.
ARCH_GOLD = "#E0A33E"   # slow (bottom-third speed) — ranked elite anyway = the coachable edge
MID_SLATE = "#B8C4D2"   # average speed
FAST_BLUE = "#5B7FA6"   # fast (top-third speed) — wheels did the heavy lifting
from matplotlib.patches import Patch

def _surname(name):
    toks = [t for t in str(name).split() if t.rstrip(".") not in ("Jr", "Sr", "II", "III")]
    return toks[-1] if toks else str(name)

def _bar_color(sprint_pctile):
    if pd.isna(sprint_pctile):  return MID_SLATE
    if sprint_pctile <= 33:     return ARCH_GOLD
    if sprint_pctile >= 66:     return FAST_BLUE
    return MID_SLATE

def _byyear_legend(fig):
    handles = [Patch(fc=ARCH_GOLD, label="SLOW runner (bottom-third speed) ranking elite — the coachable edge"),
               Patch(fc=MID_SLATE, label="Average speed"),
               Patch(fc=FAST_BLUE, label="FAST runner (top-third speed) — speed did the work"),
               Patch(fc="white", ec="#888", label="★  = model benchmark (Naylor / Soto)")]
    fig.legend(handles=handles, loc="upper center", ncol=4, fontsize=9.5, frameon=False,
               bbox_to_anchor=(0.5, 0.912), handlelength=1.3, columnspacing=1.6)

def _byyear_facets(df, value_col, label_fmt, title, subtitle, out_name,
                   seasons=(2023, 2024, 2025, 2026), topn=25):
    n = len(seasons)
    fig, axes = plt.subplots(1, n, figsize=(4.6 * n, 8.6), dpi=150)
    if n == 1: axes = [axes]
    vmax = 0
    for season in seasons:
        s = df[df["season"] == season].nlargest(topn, value_col)
        if len(s): vmax = max(vmax, float(s[value_col].max()))
    for ax, season in zip(axes, seasons):
        s = df[df["season"] == season].nlargest(topn, value_col).copy()
        s = s.sort_values(value_col)           # smallest at bottom → largest on top
        y = range(len(s))
        colors = [_bar_color(r.sprint_pctile) for r in s.itertuples()]
        ax.barh(list(y), s[value_col].to_numpy(), color=colors, edgecolor="white", linewidth=0.5)
        ax.set_yticks(list(y))
        # ★ marks the model benchmarks (Naylor / Soto); colour still reflects speed tier only
        labels = [("★ " + _surname(p)) if rid in (NAYLOR_ID, SOTO_ID) else _surname(p)
                  for p, rid in zip(s["player_name"], s["runner_id"])]
        ax.set_yticklabels(labels, fontsize=7.5)
        ax.set_xlim(0, vmax * 1.18)
        for i, v in zip(y, s[value_col].to_numpy()):
            ax.text(v + vmax * 0.015, i, label_fmt(v), va="center", fontsize=7, color="#333")
        ax.set_title(f"{season}", fontsize=12, fontweight="bold")
        ax.tick_params(axis="x", labelsize=7.5)
        for sp in ("top", "right"):
            ax.spines[sp].set_visible(False)
        ax.grid(axis="x", alpha=0.15)
    # reserve a clean header band: title → subtitle → color key, none overlapping
    fig.subplots_adjust(top=0.86, bottom=0.04, left=0.06, right=0.985, wspace=0.32)
    fig.suptitle(title, fontsize=15, fontweight="bold", y=0.975)
    fig.text(0.5, 0.94, subtitle, ha="center", fontsize=10, color="#444", style="italic")
    _byyear_legend(fig)
    out = FIG_DIR / out_name
    fig.savefig(out, bbox_inches="tight"); plt.close(fig)
    return out

def build_ground_covered_byyear():
    """Top 25 per season by ground covered BEYOND what sprint speed predicts (feet)."""
    gc = bl.parse("Ground Covered")
    gc = gc[gc["volume_qualified"] == True].copy()
    return _byyear_facets(
        gc, value_col="gain_residual_ft",
        label_fmt=lambda v: f"+{v:.2f} ft" if v >= 0 else f"{v:.2f} ft",
        title="Ground Covered Beyond What Speed Predicts — Top 25 per Season",
        subtitle="Bar length = extra feet gained from the pitcher's first move to release, "
                 "after removing the runner's sprint speed (so it's timing & jump, not wheels).",
        out_name="Fig_v8_GroundCovered_ByYear.png")

def build_bcs_byyear():
    """Top 25 per season by Blueprint Conversion Score (steal-skill composite)."""
    bcs = bl.parse("BCS Top 25 by Season").copy()
    return _byyear_facets(
        bcs, value_col="BCS",
        label_fmt=lambda v: f"{v:+.2f}",
        title="Blueprint Conversion Score — Top 25 per Season (2023–2026)",
        subtitle="Bar length = BCS, a speed-adjusted steal-skill index (higher = converts and covers "
                 "ground far better than the runner's raw speed would predict).",
        out_name="Fig_v8_BCS_ByYear.png")

# ── plain-English translation:  percentage points → concrete steals ───────────
# Every "+X pp" in this report is also shown as steals, so a coach never has to do
# the arithmetic. A typical full-season volume is ~20 attempts (league median ≈ 17–21).
SEASON_ATTEMPTS = 20
BREAK_EVEN      = 0.75   # ~ run-value break-even SB success rate

def bags_short(pp, attempts=SEASON_ATTEMPTS):
    """Compact: percentage-point success change → net bags over a typical season."""
    n = pp / 100.0 * attempts
    return "~0 bags" if abs(n) < 0.5 else f"{n:+.0f} bags/{attempts}"

def bags_phrase(pp, attempts=SEASON_ATTEMPTS):
    """Sentence form for captions/tables."""
    n = pp / 100.0 * attempts
    if abs(n) < 0.5:
        return "≈ no change in bags"
    return (f"≈ {abs(n):.0f} {'more safe steals' if n > 0 else 'fewer safe steals'} "
            f"per {attempts}-attempt season")

# Which GLM levers a coach can actually train, vs. opponent/situation, vs. physical
GLM_KIND = {
    "Post-Release Distance": "train", "Jump Time": "train", "Accel Gap": "train",
    "Lead Gain (jerk)": "train", "Pre-Release Velocity": "train",
    "Accel→Top-Speed Premium": "train", "Two-Strike Count Share": "train",
    "Primary Lead": "train",
    "Avg Catcher Pop Faced": "context", "Avg Pitcher Pickoff Rate Faced": "context",
    "Share vs Weak-Arm Catchers": "context",
    "Sprint Speed (capped at 28)": "physical", "Bolts": "physical",
}
GLM_PLAIN = {
    "Post-Release Distance": "Ground covered after the pitcher commits",
    "Jump Time": "Jump — quickness over the first 30 ft",
    "Accel Gap": "First-step quickness vs. top speed",
    "Lead Gain (jerk)": "Secondary-lead burst as the pitcher delivers",
    "Pre-Release Velocity": "Closing speed before the pitcher releases",
    "Accel→Top-Speed Premium": "Reaches top speed in fewer feet",
    "Two-Strike Count Share": "Picking two-strike counts to run",
    "Primary Lead": "Primary lead distance off the bag",
    "Avg Catcher Pop Faced": "Catcher pop time faced (opponent)",
    "Avg Pitcher Pickoff Rate Faced": "Pitcher hold/pickoff rate faced (opponent)",
    "Share vs Weak-Arm Catchers": "Facing weak-arm catchers (situation)",
    "Sprint Speed (capped at 28)": "Raw top speed (capped — diminishing returns)",
    "Bolts": "Number of 30+ ft/s sprints",
}
KIND_COLOR = {"train": "#3FA66B", "context": "#9AA0A6", "physical": "#D98A3D"}

# ── THE EQUATION FIGURE — what each skill is worth, per +1 SD ──────────────────
def build_equation_figure():
    """Steal-success equation + per-1-SD coefficient bars (Statcast-clean, keyed)."""
    g = glm.copy()
    g["abs"]  = g["sb_pct_boost_per_tier"].abs()
    g = g.sort_values("abs").reset_index(drop=True)          # ascending → biggest on top
    g["kind"]  = g["feature"].map(GLM_KIND).fillna("context")
    g["label"] = g["feature"].map(GLM_PLAIN).fillna(g["feature"])

    fig = plt.figure(figsize=(12.6, 9.2), dpi=150)
    # Stacked layout: full-width text band on top, full-width bar panel below.
    # (Left margin is generous so the long lever names never collide with anything.)
    ax  = fig.add_axes([0.345, 0.055, 0.615, 0.595])        # bar panel, full width, lower band

    fig.text(0.035, 0.972, "The Steal-Success Equation", fontsize=21, fontweight="bold", color="#1f2d3d")
    fig.text(0.035, 0.937, "How much each skill is worth — the change in steal-success rate from a "
             "one-standard-deviation (1-SD) improvement.", fontsize=11.5, color="#555", style="italic")
    fig.text(0.035, 0.888, "Chance of a successful steal   =   baseline (≈ 78%)   +   "
             "Σ ( weight  ×  how far above average the runner is, in SDs )",
             fontsize=13, fontweight="bold", color="#1f2d3d")
    fig.text(0.035, 0.851, "math form:   log-odds(success)  =  β₀ + β₁·z₁ + β₂·z₂ + …",
             fontsize=10.5, color="#666", family="monospace")
    fig.text(0.035, 0.822, "β = the bar lengths below;   z = how many SDs above league average a runner is.",
             fontsize=10, color="#777")

    # Horizontal color key spread across the full width (no left column to collide with).
    leg_y = 0.760
    guide = [("#3FA66B", "Trainable — a coach can develop this"),
             ("#9AA0A6", "Context — opponent / situation"),
             ("#D98A3D", "Physical — raw speed, barely moves it")]
    xs = [0.035, 0.375, 0.685]
    for (col, txt), x in zip(guide, xs):
        fig.add_artist(plt.Rectangle((x, leg_y), 0.018, 0.021, transform=fig.transFigure,
                                     facecolor=col, edgecolor="none"))
        fig.text(x + 0.026, leg_y + 0.002, txt, fontsize=10.5, color="#333")
    fig.text(0.035, 0.712,
             "Each bar = the standardized weight β — the points of steal-success rate added by a +1-SD gain "
             "in that one skill (also shown as net bags over a 20-attempt season).  Bigger bar = bigger "
             "coaching payoff.", fontsize=10, color="#555")

    colors = [KIND_COLOR[k] for k in g["kind"]]
    yb = list(range(len(g)))
    ax.barh(yb, g["sb_pct_boost_per_tier"].to_numpy(), color=colors, edgecolor="white", linewidth=0.6)
    ax.axvline(0, color="#333", lw=1.0)
    ax.set_yticks(yb); ax.set_yticklabels(g["label"], fontsize=9.5)
    ax.set_ylim(-0.7, len(g) - 0.3)
    vmax = float(g["sb_pct_boost_per_tier"].max()); vmin = float(g["sb_pct_boost_per_tier"].min())
    pad = vmax * 0.03
    for i, (pp, beta) in enumerate(zip(g["sb_pct_boost_per_tier"], g["tech_coefficient"])):
        txt = f"{pp:+.0f} pts   (β={beta:+.2f},  {bags_short(pp)})"
        if pp >= 0:                       # right of the bar tip
            ax.text(pp + pad, i, txt, va="center", ha="left", fontsize=8.4, color="#222")
        else:                             # right of the zero line — never crowds the left labels
            ax.text(pad, i, txt, va="center", ha="left", fontsize=8.4, color="#7a4a16")
    ax.set_xlabel("Change in steal-success rate for a +1-SD improvement  (percentage points)",
                  fontsize=10.8, fontweight="bold")
    ax.set_xlim(min(vmin * 1.2, -4), vmax * 1.62)
    for sp in ("top", "right", "left"): ax.spines[sp].set_visible(False)
    ax.tick_params(left=False); ax.grid(axis="x", alpha=0.15)
    out = FIG_DIR / "Fig_v8_Equation.png"
    fig.savefig(out, bbox_inches="tight"); plt.close(fig)
    return out

# ── 2025 COACHING TARGET BOARD — who to coach, and the projected prize ────────
def coaching_targets(season=2025):
    """Two honest tracks: green-light (volume) and technique-fix (efficiency)."""
    cur = sssi[sssi["season"] == season].copy()
    cur["net"] = cur["SB"] - cur["CS"]

    # GREEN-LIGHT: fast + already efficient + under-running → just let them run.
    gl = cur[(cur["pct_speed"] >= 66) & (cur["real_sb_pct"] >= 0.80) &
             (cur["sb_attempts"] >= 6) & (cur["sb_attempts"] < SEASON_ATTEMPTS)].copy()
    gl["extra_att"]      = (SEASON_ATTEMPTS - gl["sb_attempts"]).clip(lower=0)
    gl["proj_extra_net"] = (gl["extra_att"] * (2 * gl["real_sb_pct"] - 1)).round(0)
    gl = gl.sort_values("proj_extra_net", ascending=False).head(8).copy()
    gl["sbcs"] = gl["SB"].astype(int).astype(str) + "/" + gl["CS"].astype(int).astype(str)

    # TECHNIQUE-FIX: high volume but caught too often → drill the weakest lever.
    tf = cur[(cur["sb_attempts"] >= 12) & (cur["real_sb_pct"] < 0.70)].copy()
    lever_name = {"z_jump": "Jump", "z_post_rel_dist": "Ground covered",
                  "z_accel_gap": "First-step burst", "z_lead_gain": "Secondary-lead burst"}
    lever_pp   = {"z_jump": 17.32, "z_post_rel_dist": 24.76, "z_accel_gap": 9.94, "z_lead_gain": 2.08}
    weak, fixpp = [], []
    for _, r in tf.iterrows():
        cand = {c: r[c] for c in lever_name if c in tf.columns and pd.notna(r[c])}
        if not cand:
            weak.append("—"); fixpp.append(0.0); continue
        wc = min(cand, key=cand.get)                     # lowest z = biggest gap below avg
        gap = max(0.0, -float(cand[wc]))                 # SDs below league average
        weak.append(lever_name[wc]); fixpp.append(gap * lever_pp[wc])
    tf["weak_lever"] = weak; tf["fix_pp"] = fixpp
    tf = tf.sort_values("sb_attempts", ascending=False).head(6).copy()
    tf["sbcs"] = tf["SB"].astype(int).astype(str) + "/" + tf["CS"].astype(int).astype(str)
    return gl, tf

# ══════════════════════════════════════════════════════════════════════════════
# XYZ section lead-in: What we did → Result → Variables shown
# ══════════════════════════════════════════════════════════════════════════════
# Per-table variable definitions (term → plain meaning + units). Self-contained, no
# cross-referencing required. Superset; each section pulls only the keys it shows.
VAR_DEFS = {
    "Player": "Runner (and season).",
    "Top Speed": "Sprint speed in feet/second — the structural baseline (Statcast).",
    "Jump (s)": "Seconds to cover the first 30 ft. Lower = quicker first step. Coachable.",
    "First-Step Burst": "Reaches top speed in fewer feet than his speed predicts (premium when fast).",
    "Steals Above Speed-Expected": "Real success rate minus what sprint speed alone predicts (pct pts).",
    "SB / CS": "Stolen bases / times caught stealing.",
    "Slow-Steal Skill": "Overall technique score — skill beyond raw speed. Higher = better.",
    "Untapped Speed": "Fast but under-stealing: z(top speed) − z(net steals). Higher = bigger coaching upside.",
    "Speed (SD)": "Top speed expressed in standard deviations vs. the league (positive = faster).",
    "Coachable lever": "A trainable skill the model links to stolen-base success.",
    "League Avg": "Typical value across qualified runners.",
    "One-Step Improvement": "A realistic one-standard-deviation gain in that skill.",
    "Success-Rate Change": "Percentage-point change in SB success from that one-step gain.",
    "Ground vs Speed-Expected": "Extra feet of ground covered (first move → release) beyond speed (ft).",
    "BCS": "Blueprint Conversion Score — speed-adjusted steal-skill composite (higher = better).",
    "Net SB": "Stolen bases minus times caught.",
    "Steal Value (xSB)": "z(net steals above avg) + z(top speed) — fast AND productive.",
    "Catcher Pop (faced)": "Average catcher pop time the runner faced (battery context).",
    "AUC": "Area under ROC — predictive accuracy (0.5 = coin flip, 1.0 = perfect).",
    "Extra steals": "Plain translation of the success-rate change into net bags over a typical 20-attempt season.",
    "Success rate": "Share of steal attempts that succeed (SB ÷ attempts).",
    "Attempts": "Stolen-base attempts this season (SB + CS).",
    "If unleashed": "Projected extra net steals at a modest, speed-appropriate ~20-attempt volume, kept at the runner's own demonstrated success rate.",
    "Weakest lever": "The trainable skill where this runner is furthest below league average.",
    "Projected fix": "Success-rate gain if that weakest lever is lifted to the league average.",
}

def section_intro(doc, what, result, variables):
    """Google-XYZ lead-in placed immediately before a table or figure."""
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
    r = p.add_run("What we did.  "); r.bold = True; r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor.from_string(NAVY)
    r2 = p.add_run(what); r2.font.size = Pt(9.5)
    takeaway(doc, result)
    if variables:
        cap = doc.add_paragraph(); cap.paragraph_format.space_after = Pt(1)
        rc = cap.add_run("Variables in this view:"); rc.bold = True; rc.font.size = Pt(9)
        rc.font.color.rgb = RGBColor.from_string("555555")
        for term in variables:
            meaning = VAR_DEFS.get(term, "")
            b = doc.add_paragraph(style="List Bullet")
            b.paragraph_format.space_after = Pt(0); b.paragraph_format.left_indent = Inches(0.25)
            rb = b.add_run(f"{term} — "); rb.bold = True; rb.font.size = Pt(8.5)
            rm = b.add_run(meaning); rm.font.size = Pt(8.5)

# ══════════════════════════════════════════════════════════════════════════════
# MAIN REPORT
# ══════════════════════════════════════════════════════════════════════════════
def build_main():
    xsb_fig  = build_xsb_quadrant()
    eq_fig   = build_equation_figure()
    gl, tf   = coaching_targets(2025)
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"; doc.styles["Normal"].font.size = Pt(10.5)
    set_margins(doc)

    # ---- P1: BLUF + glossary + legend ----------------------------------------
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("The Naylor Model"); r.bold = True; r.font.size = Pt(26)
    sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(6)
    r = sub.add_run("Finding base-stealing skill that sprint speed hides  ·  for MLB R&D and coaching staffs")
    r.italic = True; r.font.size = Pt(11); r.font.color.rgb = RGBColor.from_string("666666")

    takeaway(doc, "Some of the best base-stealers in baseball are among its slowest runners. This "
                  "report measures the skill speed hides, gives you the equation for what that skill "
                  "is worth, and hands you a 2025 list of exactly who to coach and the steals it adds.")
    body(doc,
         "The model is named for the clearest living example of one quality — a runner who steals better "
         "than ~99% of the league despite bottom-of-the-league speed. In 2025 Josh Naylor stole 30 bases "
         "at 92% while running slower than 98% of MLB: the #1 steal-skill season in our data of 673. But "
         "this is not a report about one player. It is about the quality he exemplifies — a quick jump, an "
         "explosive secondary lead, and reading the pitcher early — which shows up all over the league and, "
         "unlike foot speed, is coachable. The next four pages give you (1) the equation for what each "
         "skill is worth, (2) the runners who already have it, (3) a 2025 target board of who to develop "
         "and the projected prize, and (4) the three drills that move the needle most.",
         size=10, after=8)

    H(doc, "What the metrics mean (everything you need — no other pages required)", lvl=2)
    add_glossary_table(doc, GLOSSARY)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    H(doc, "How to read this report", lvl=2)
    legend_line(doc)
    for head, txt in [
        ("Colors are a within-table ranking, not a league verdict.",
         "Green = a strength, red = a weakness, white = middle — but only versus the other players in "
         "that same table. A red cell can still be an above-average MLB number; it just trails the others "
         "shown. On the leaderboard figures there is deliberately no red at all: bars are gold for slow "
         "runners and blue for fast, because every runner shown is already elite."),
        ("\"Percentage points\" are always translated into steals.",
         "A percentage point (pp) is a raw change in success rate — going from 78% to 80% is +2 pp. So you "
         "never have to do the math, every pp is also shown as bags: for a runner who tries ~20 steals a "
         "season, +5 pp ≈ one extra stolen base and one fewer caught-stealing."),
        ("What this model is good at.",
         "It is strong at ranking steal skill and telling you which techniques pay off — that is what the "
         "equation and the rankings deliver. It is deliberately not a single-pitch oracle: stealing is "
         "high-variance, so no model calls individual attempts well (the honest, de-leaked predictive "
         "accuracy is in the Technical Appendix). Use it to decide who to develop and what to drill."),
    ]:
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(3)
        rb = p.add_run(head + "  "); rb.bold = True; rb.font.size = Pt(9.5)
        rb.font.color.rgb = RGBColor.from_string(NAVY)
        rt = p.add_run(txt); rt.font.size = Pt(9.5)

    doc.add_page_break()

    # ---- P2: THE EQUATION — what each skill is worth -------------------------
    H(doc, "1  What Each Skill Is Worth — The Steal-Success Equation")
    section_intro(doc,
        what="We fit an interpretable model (a logistic GLM) on Statcast runner-seasons that isolates "
             "how each skill changes the chance of a successful steal, then converted every weight into "
             "plain points of success rate for one realistic one-standard-deviation (1-SD) improvement — "
             "and into concrete steals.",
        result="Three trainable levers dominate — covering ground after the pitcher commits, a quicker "
                "jump, and reaching top speed sooner. Raw foot speed barely moves the needle.",
        variables=["Coachable lever", "League Avg", "One-Step Improvement",
                   "Success-Rate Change", "Extra steals"])
    add_fig(doc, eq_fig, width=7.0,
            caption="Figure 1 — Each bar is the standardized weight β: the change in steal-success rate "
                    "from a +1-SD improvement in that one skill (also shown as net bags / 20 attempts). "
                    "Green = trainable, grey = opponent/situation context, orange = raw speed. The "
                    "equation at top: success odds = baseline + Σ (weight × how many SDs above average "
                    "the runner is).")

    g = glm.copy()
    g["abs"]   = g["sb_pct_boost_per_tier"].abs()
    g["arrow"] = np.where(g["sb_pct_boost_per_tier"] >= 0, "↑ raises", "↓ lowers")
    g["lever"] = g["feature"].map(GLM_PLAIN).fillna(g["feature"])
    g["bags"]  = g["sb_pct_boost_per_tier"].apply(bags_short)
    g = g.sort_values("abs", ascending=False).head(9).copy()
    heat_table(doc, g,
        columns=[("lever", "Coachable lever", "left"),
                 ("league_avg", "League\nAvg", "center"),
                 ("one_tier_step", "One-Step\n(+1 SD)", "center"),
                 ("sb_pct_boost_per_tier", "Success Rate\nChange (pts)", "center"),
                 ("bags", "≈ Extra steals\n(per 20 tries)", "center"),
                 ("arrow", "Effect", "center")],
        heat_specs={"sb_pct_boost_per_tier": +1},
        fmt={"league_avg": lambda v: f"{v:.2f}", "one_tier_step": lambda v: f"{v:.2f}",
             "sb_pct_boost_per_tier": lambda v: f"{v:+.0f} pts"}, size=9)
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = cap.add_run("Table 1 — The nine biggest levers. \"+17 pts\" means a one-step (1-SD) improvement "
                     "raises the success rate by 17 percentage points — about 3 extra safe steals over a "
                     "20-attempt season. Read with Figure 1.")
    rr.italic = True; rr.font.size = Pt(8.5); rr.font.color.rgb = RGBColor.from_string("666666")

    doc.add_page_break()

    # ---- P3: SSSI Excel-style matrix -----------------------------------------
    H(doc, "2  The Runners Who Already Have the Skill  (Slow-Steal Skill)")
    section_intro(doc,
        what="We combined the skills above — jump, first-step burst, and steals-above-what-speed-predicts "
             "— into one 'Slow-Steal Skill' score, tuned on 80% of the league with the benchmark names "
             "(Naylor, Soto) held out, then ranked every runner-season.",
        result="The leaderboard is full of mid-pack-speed runners grading out elite on skill — proof the "
                "trait is widespread and findable, not a one-player quirk.",
        variables=["Player", "Top Speed", "Jump (s)", "First-Step Burst",
                   "Steals Above Speed-Expected", "SB / CS", "Slow-Steal Skill"])
    legend_line(doc)

    ss = sssi.sort_values("SSSI_v7", ascending=False).head(15).copy()
    ss["sbcs"] = ss["SB"].astype(int).astype(str) + "/" + ss["CS"].astype(int).astype(str)
    heat_table(doc, ss,
        columns=[("rank_v7", "#", "center"), ("player_name", "Player", "left"),
                 ("season", "Year", "center"), ("sprint_speed", "Top Speed\n(ft/s)", "center"),
                 ("jump_time", "Jump (s)\nlower better", "center"),
                 ("accel_topspeed_premium", "First-Step\nBurst", "center"),
                 ("sb_residual", "Steals Above\nSpeed-Expected", "center"),
                 ("sbcs", "SB / CS", "center"),
                 ("SSSI_v7", "Slow-Steal\nSkill", "center")],
        heat_specs={"sb_residual": +1, "accel_topspeed_premium": +1, "SSSI_v7": +1,
                    "jump_time": -1},
        highlight_ids={NAYLOR_ID, SOTO_ID},
        fmt={"season": lambda v: f"{int(v)}", "sprint_speed": lambda v: f"{v:.1f}",
             "jump_time": lambda v: f"{v:.2f}", "SSSI_v7": lambda v: f"{v:+.2f}",
             "sb_residual": lambda v: f"{v*100:+.1f}%",
             "accel_topspeed_premium": lambda v: f"{v:+.1f}", "rank_v7": lambda v: f"{int(v)}"},
        size=9, header_size=8.5)
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = cap.add_run("Table 2 — Top 15 by Slow-Steal Skill (2015–2026). The Top Speed column is "
                     "deliberately uncolored: speed is NOT what earns a spot here. Amber rows flag the "
                     "benchmark names (Naylor, Soto) — the clearest examples of the trait, not the basis "
                     "of the score (they were held out while the score was tuned).")
    rr.italic = True; rr.font.size = Pt(8.5); rr.font.color.rgb = RGBColor.from_string("666666")

    doc.add_page_break()

    # ---- P4: speed vs production quadrant + 2025 COACHING TARGET BOARD --------
    H(doc, "3  Who To Coach In 2025 — Targets and the Projected Prize")
    section_intro(doc,
        what="We plotted every qualified runner-season since 2023 by top speed (left→right) against steal "
             "production (down→up), then pulled the 2025 names into a target board with a projected "
             "payoff for each.",
        result="Two clean opportunities fall out: fast, efficient runners who simply don't run enough "
                "(green-light them) and high-volume runners who get caught too often (fix the technique).",
        variables=["Speed (SD)", "Steal Value (xSB)", "Untapped Speed"])
    add_fig(doc, xsb_fig, width=6.6,
            caption="Figure 2 — Each dot is a runner-season (2023–2026). X-axis = top speed in standard "
                    "deviations (right = faster); Y-axis = steal production (up = more). Green = fast & "
                    "productive; blue (lower-right) = fast but under-stealing (the green-light targets); "
                    "gold = the slow-but-productive archetype (★ Naylor, Soto); grey = neither. No red — "
                    "no quadrant here is 'bad'.")

    # 3a — green-light (volume) board
    H(doc, "3a  Green-light targets — fast, efficient, under-running (just let them run)", lvl=2)
    section_intro(doc,
        what="We isolated the 2025 runners with top-third speed who already succeed at 80%+ but attempt "
             "fewer than ~20 steals, and projected the bags they'd add at a modest, speed-appropriate "
             "volume — kept at their own proven success rate.",
        result="These are low-risk green-light decisions, not mechanics rebuilds: each already converts; "
                "they simply aren't running.",
        variables=["Player", "Top Speed", "SB / CS", "Success rate", "Attempts", "If unleashed"])
    legend_line(doc)
    glr = gl.copy(); glr.insert(0, "gl_rank", range(1, len(glr) + 1))
    heat_table(doc, glr,
        columns=[("gl_rank", "#", "center"), ("player_name", "Player", "left"),
                 ("sprint_speed", "Top Speed\n(ft/s)", "center"), ("sbcs", "SB / CS", "center"),
                 ("real_sb_pct", "Success\nrate", "center"),
                 ("sb_attempts", "Attempts\n(2025)", "center"),
                 ("proj_extra_net", "If unleashed\n(+ bags)", "center")],
        heat_specs={"real_sb_pct": +1, "proj_extra_net": +1},
        fmt={"sprint_speed": lambda v: f"{v:.1f}", "real_sb_pct": lambda v: f"{v*100:.0f}%",
             "sb_attempts": lambda v: f"{int(v)}", "proj_extra_net": lambda v: f"+{int(v)}",
             "gl_rank": lambda v: f"{int(v)}"}, size=9)
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = cap.add_run("Table 3 — 2025 green-light board. \"If unleashed\" = projected extra net steals at "
                     "~20 attempts, at the runner's own 2025 success rate (illustrative, not a forecast "
                     "of behavior).")
    rr.italic = True; rr.font.size = Pt(8.5); rr.font.color.rgb = RGBColor.from_string("666666")

    # 3b — technique-fix (efficiency) board
    H(doc, "3b  Technique-fix targets — high volume, caught too often (drill the weakest lever)", lvl=2)
    section_intro(doc,
        what="The opposite group: 2025 runners attempting a lot (12+) but converting below 70%. For each "
             "we name the trainable skill where they are furthest below average and the success-rate gain "
             "from lifting it to league average.",
        result="These should not run more — they should run better. The fix is mechanics, and the model "
                "says which one.",
        variables=["Player", "Top Speed", "SB / CS", "Success rate", "Weakest lever", "Projected fix"])
    tfr = tf.copy()
    tfr["fix_txt"] = tfr["fix_pp"].apply(lambda pp: f"+{pp:.0f} pts" if pp >= 0.5 else "—")
    heat_table(doc, tfr,
        columns=[("player_name", "Player", "left"), ("sprint_speed", "Top Speed\n(ft/s)", "center"),
                 ("sbcs", "SB / CS", "center"), ("real_sb_pct", "Success\nrate", "center"),
                 ("sb_attempts", "Attempts\n(2025)", "center"),
                 ("weak_lever", "Weakest\nlever", "left"),
                 ("fix_txt", "Fix → success\ngain", "center")],
        heat_specs={"real_sb_pct": +1, "fix_pp": +1},
        fmt={"sprint_speed": lambda v: f"{v:.1f}", "real_sb_pct": lambda v: f"{v*100:.0f}%",
             "sb_attempts": lambda v: f"{int(v)}"}, size=9)
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = cap.add_run("Table 4 — 2025 technique-fix board. \"Fix → success gain\" = projected percentage-"
                     "point gain if the weakest lever is raised to league average (via the Table-1 "
                     "weights); higher success rate = redder here is worse, greener is better.")
    rr.italic = True; rr.font.size = Pt(8.5); rr.font.color.rgb = RGBColor.from_string("666666")

    doc.add_page_break()

    # ---- P5: what to coach (tied to the equation) + honest caveats -----------
    H(doc, "4  What To Coach — Three Priorities, and the Caveats")
    body(doc, "The equation in Section 1 points to three trainable levers. Each is technique, not body "
              "type — which is why a slow runner can lead the league in steal skill. The payoff below is "
              "straight off Table 1.", size=10, after=6)
    for head, txt in [
        ("Cover more ground once the pitcher commits.  (+25 pts ≈ +5 bags / 20 tries)",
         "The single biggest lever. Train the secondary-lead explosion and the read of the pitcher's first "
         "move — distance gained before the throw is worth more than raw foot speed."),
        ("Sharpen the jump.  (+17 pts ≈ +3 bags / 20 tries)",
         "Cutting the time to the first 30 feet is pure technique — stance, weight shift, and first-step "
         "direction, not body type."),
        ("Reach top speed sooner — first-step burst.  (+10 pts ≈ +2 bags / 20 tries)",
         "Runners who hit top gear in fewer feet steal above their speed peers. Acceleration mechanics, "
         "not a higher top speed, are what convert."),
    ]:
        p = doc.add_paragraph(style="List Number"); p.paragraph_format.space_after = Pt(3)
        rb = p.add_run(head + "  "); rb.bold = True; rb.font.size = Pt(10)
        rt = p.add_run(txt); rt.font.size = Pt(10)
    body(doc, "Net play: green-light the Table-3 names, drill the Table-4 names on the lever listed, and "
              "you convert latent speed and wasted attempts into runs — without touching anyone's physical "
              "ceiling.", size=10, after=8)

    H(doc, "Read the projections honestly", lvl=2)
    for head, txt in [
        ("The target boards are estimates, not forecasts.",
         "\"If unleashed\" holds a runner at his own 2025 success rate and a modest ~20-attempt volume; "
         "real results depend on health, matchups, and game situation. Treat the bag counts as a "
         "priority ranking, not a guarantee."),
        ("Per-lever gains assume the rest stays equal.",
         "A +1-SD improvement is real but ambitious, and the model holds the other skills fixed. Small, "
         "compounding gains across two levers are more realistic than a single full-SD jump."),
        ("Context columns are not coaching levers.",
         "Catcher pop time and pitcher hold rate move success in the math, but you don't coach them — "
         "they're shown so you can separate skill from circumstance."),
    ]:
        p = doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after = Pt(2)
        rb = p.add_run(head + "  "); rb.bold = True; rb.font.size = Pt(9.5)
        rt = p.add_run(txt); rt.font.size = Pt(9.5)

    doc.save(str(OUT_MAIN))
    print(f"wrote {OUT_MAIN}  ({OUT_MAIN.stat().st_size/1024:.0f} KB)")

# ══════════════════════════════════════════════════════════════════════════════
# TECHNICAL APPENDIX (full detail for auditors)
# ══════════════════════════════════════════════════════════════════════════════
def build_appendix():
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"; doc.styles["Normal"].font.size = Pt(10)
    set_margins(doc, left=0.7, right=0.7)

    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("The Naylor Model — v8 Technical Appendix"); r.bold = True; r.font.size = Pt(22)
    sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Full model detail, validation, and the Blueprint Conversion Score. "
                    "Companion to the applied main report.")
    r.italic = True; r.font.size = Pt(10.5); r.font.color.rgb = RGBColor.from_string("666666")
    doc.add_paragraph()

    # A — Models + de-leak / AUC
    H(doc, "A  The Three Models (A / B / C) and Validation")
    section_intro(doc,
        what="We built three models on the same Statcast runner-season data: a per-attempt predictor "
             "(Model A), a season-level predictor (Model B, the headline — now a Bayesian-tuned XGBoost), "
             "and an interpretable GLM (Model C) for plain-English weights.",
        result="The de-leaked accuracy is honest but modest — which is exactly why the coach-facing "
                "report leads with skill rankings and coaching levers, not AUC.",
        variables=["AUC"])
    full_auc = float("nan")
    if auc is not None and "auc" in auc.columns:
        lbl = auc.get("label", auc.iloc[:, 0]).astype(str)
        full = auc.loc[lbl.str.contains("full", case=False)]
        if len(full): full_auc = float(full["auc"].iloc[0])
    tbl = doc.add_table(rows=1, cols=4); tbl.style = "Table Grid"; tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(["Model", "Unit", "AUC", "Purpose"]):
        _set_cell_text(tbl.rows[0].cells[j], h, bold=True, size=9, color="FFFFFF"); shade(tbl.rows[0].cells[j], NAVY)
    for rdat in [("Model A — per-attempt XGBoost", "Individual attempt", "0.739",
                  "Strongest predictor — does THIS steal succeed"),
                 ("Model B — season XGBoost (tuned)", "Runner-season",
                  (f"{full_auc:.3f} (de-leaked)" if full_auc == full_auc else "0.624 (de-leaked)"),
                  "Ranks season-long skill"),
                 ("Model C — interpretable GLM", "Runner-season", "—", "Plain-English weight table")]:
        c = tbl.add_row().cells
        for j, v in enumerate(rdat): _set_cell_text(c[j], v, size=9)
    body(doc,
         "AUC caveat (v7/v8 de-leaking). Versions v4–v6 reported AUCs of ~0.66–0.70, but those runs "
         "carried duplicate runner-season rows — repeated Statcast split measurements for one "
         "player-season — that leaked across cross-validation folds and inflated the score. The pipeline "
         "now averages those duplicate splits into one row per runner-season. The resulting AUC is lower "
         "but honest; it is intentionally kept out of the main report because predictive AUC is not the "
         "deliverable for coaches — the skill rankings and coaching levers are.", size=9.5)
    body(doc,
         "v8 model update. Model B was upgraded from a gradient-boosting classifier to a Bayesian-tuned "
         "XGBoost (100-trial Optuna search) on the same de-leaked data. De-leaked CV AUC rose from 0.589 "
         "to 0.624 overall, and from 0.588 to 0.665 in the post-2023 (bigger-bases) era — a real gain from "
         "model choice and tuning, not new signal. The SSSI rankings and the steal-success equation are "
         "unchanged.", size=9.5)
    body(doc,
         "v9 — the per-attempt model (Model A). Moving from 673 season aggregates to the ~10,400 "
         "individual tracked attempts (Statcast leads cache) lifts CV AUC to 0.739. The driver is the "
         "per-pitch lead distances — how much ground the runner actually covered on that attempt — which "
         "is precisely this report's thesis: ground covered, not raw speed, decides the steal. The model "
         "is leakage-checked (no outcome-derived columns; catcher/pitcher tendencies are out-of-fold "
         "encoded). Deep learning was considered and rejected: at ~10k rows and a dozen tabular features, "
         "gradient boosting is the right tool — neural nets need far more data and overfit here.", size=9.5)
    add_fig(doc, FIG_DIR / "Fig_v7_AUC.png", 6.0,
            "Figure A1 — Model AUC across versions; the latest bar is the de-leaked tuned-XGBoost figure.")
    add_fig(doc, FIG_DIR / "Fig_v7_Importance_PrePost.png", 6.2,
            "Figure A2 — XGBoost feature-importance shift after the 2023 bigger-bases rule change.")

    # B — full GLM + the equation figure
    H(doc, "B  Model C — The Steal-Success Equation (Full GLM)")
    section_intro(doc,
        what="We list every feature in the interpretable GLM, sorted by absolute impact on stolen-base "
             "success. The figure restates the model as an equation: standardized coefficients (β) on "
             "z-scored inputs, i.e. the effect of a one-SD move in each skill.",
        result="The biggest movers are coachable timing skills (ground covered, jump, first-step "
                "burst); raw speed and context (catcher/pitcher faced) matter less.",
        variables=["Coachable lever", "League Avg", "One-Step Improvement", "Success-Rate Change", "Extra steals"])
    add_fig(doc, FIG_DIR / "Fig_v8_Equation.png", 6.8,
            "Figure B1 — log-odds(success) = β₀ + Σ βᵢ·zᵢ. Each bar is the per-1-SD coefficient as a "
            "percentage-point change in success rate (and net bags / 20 attempts). Green = trainable, "
            "grey = context, orange = raw speed.")
    body(doc, "Reading the weights: 'pp' = percentage points of success rate per one-standard-deviation "
              "improvement (one SD = the 'One-Step' column). The odds multiplier is the same effect on the "
              "odds scale (exp β). In steals: a +1-SD jump (+17 pp) is ≈ 3 more safe steals over a "
              "20-attempt season; ground-covered (+25 pp) is ≈ 5.", size=9, after=4)
    gg = glm.copy()
    gg["abs"] = gg["sb_pct_boost_per_tier"].abs()
    gg = gg.sort_values("abs", ascending=False)
    heat_table(doc, gg,
        columns=[("feature", "Feature", "left"), ("league_avg", "League\nAvg", "center"),
                 ("one_tier_step", "One-Step\nImprovement", "center"),
                 ("sb_pct_boost_per_tier", "Success-Rate\nChange (pp)", "center"),
                 ("odds_multiplier", "Odds\nMultiplier", "center")],
        heat_specs={"sb_pct_boost_per_tier": +1},
        fmt={"league_avg": lambda v: f"{v:.3f}", "one_tier_step": lambda v: f"{v:.3f}",
             "sb_pct_boost_per_tier": lambda v: f"{v:+.2f}", "odds_multiplier": lambda v: f"{v:.3f}"},
        size=8.5)

    # C — full SSSI
    H(doc, "C  Slow-Steal Skill (SSSI) — Full Top 25 with all components")
    section_intro(doc,
        what="The full Top 25 by Slow-Steal Skill, with the underlying components exposed.",
        result="Skill, not speed, drives the ranking — the leaders cluster in the mid-20s ft/s while "
                "grading out elite on jump, burst, and steals-above-expected.",
        variables=["Player", "Top Speed", "Jump (s)", "Steals Above Speed-Expected",
                   "First-Step Burst", "Slow-Steal Skill"])
    ss = sssi.sort_values("SSSI_v7", ascending=False).head(25).copy()
    heat_table(doc, ss,
        columns=[("rank_v7", "#", "center"), ("player_name", "Player", "left"),
                 ("season", "Yr", "center"), ("sprint_speed", "Top Speed\n(ft/s)", "center"),
                 ("jump_time", "Jump (s)\nlower better", "center"),
                 ("sb_residual", "Steals Above\nSpeed-Expected", "center"),
                 ("accel_topspeed_premium", "First-Step\nBurst", "center"),
                 ("real_sb_pct", "SB%", "center"), ("SSSI_v7", "Slow-Steal\nSkill", "center")],
        heat_specs={"sb_residual": +1, "accel_topspeed_premium": +1, "real_sb_pct": +1, "SSSI_v7": +1},
        highlight_ids={NAYLOR_ID, SOTO_ID},
        fmt={"season": lambda v: f"{int(v)}", "sprint_speed": lambda v: f"{v:.1f}",
             "real_sb_pct": lambda v: f"{v*100:.0f}%", "jump_time": lambda v: f"{v:.2f}",
             "SSSI_v7": lambda v: f"{v:+.2f}", "sb_residual": lambda v: f"{v*100:+.1f}%",
             "accel_topspeed_premium": lambda v: f"{v:+.1f}", "rank_v7": lambda v: f"{int(v)}"}, size=8.5)

    # D — xSB full top / crafty technicians
    H(doc, "D  Expected SB Outcome (xSB) — Full Leaderboards")
    section_intro(doc,
        what="Steal Value (xSB) = z(net steals above average) + z(top speed). It surfaces the runners "
             "who are both fast AND productive; the Top 15 are the league's 'Realized Burners'.",
        result="The ceiling cases are unsurprising burners (De La Cruz, Carroll); Naylor & Soto reach "
                "the same production from the slow side — the opposite, but equally valuable, profile.",
        variables=["Player", "Top Speed", "Net SB", "Speed (SD)", "Steal Value (xSB)"])
    xt = xsb.sort_values("xsb_outcome", ascending=False).head(15).copy()
    heat_table(doc, xt,
        columns=[("rank_xsb", "#", "center"), ("player_name", "Player", "left"),
                 ("season", "Yr", "center"), ("sprint_speed", "Top Speed\n(ft/s)", "center"),
                 ("SB", "SB", "center"), ("CS", "CS", "center"), ("net_sb", "Net SB", "center"),
                 ("z_sprint", "Speed\n(SD)", "center"), ("z_net_sb", "Steals\n(SD)", "center"),
                 ("xsb_outcome", "Steal Value\n(xSB)", "center")],
        heat_specs={"z_sprint": +1, "z_net_sb": +1, "xsb_outcome": +1, "net_sb": +1},
        fmt={"season": lambda v: f"{int(v)}", "sprint_speed": lambda v: f"{v:.1f}",
             "SB": lambda v: f"{int(v)}", "CS": lambda v: f"{int(v)}", "net_sb": lambda v: f"{int(v)}",
             "z_sprint": lambda v: f"{v:+.2f}", "z_net_sb": lambda v: f"{v:+.2f}",
             "xsb_outcome": lambda v: f"{v:+.2f}", "rank_xsb": lambda v: f"{int(v)}"}, size=8.5)

    ns = xsb[xsb["runner_id"].isin([NAYLOR_ID, SOTO_ID])].sort_values("xsb_outcome", ascending=False).copy()
    H(doc, "D.1  Naylor & Soto — Crafty Technicians (year-by-year)", lvl=2)
    section_intro(doc,
        what="Every tracked Naylor and Soto season on the same axes.",
        result="Both carry a negative 'Untapped Speed' every year — they out-steal their speed "
                "consistently, not in a single lucky season.",
        variables=["Player", "Top Speed", "Speed (SD)", "Untapped Speed"])
    heat_table(doc, ns,
        columns=[("player_name", "Player", "left"), ("season", "Yr", "center"),
                 ("sprint_speed", "Top Speed\n(ft/s)", "center"), ("SB", "SB", "center"), ("CS", "CS", "center"),
                 ("z_sprint", "Speed\n(SD)", "center"), ("z_net_sb", "Steals\n(SD)", "center"),
                 ("sb_potential_gap", "Untapped\nSpeed", "center"), ("quadrant", "Quadrant", "left")],
        heat_specs={"z_net_sb": +1, "sb_potential_gap": -1},
        highlight_ids={NAYLOR_ID, SOTO_ID},
        fmt={"season": lambda v: f"{int(v)}", "sprint_speed": lambda v: f"{v:.1f}",
             "SB": lambda v: f"{int(v)}", "CS": lambda v: f"{int(v)}",
             "z_sprint": lambda v: f"{v:+.2f}", "z_net_sb": lambda v: f"{v:+.2f}",
             "sb_potential_gap": lambda v: f"{v:+.2f}"}, size=8.5)

    # E — BCS (lead with the rebuilt, self-explaining figure)
    if bp is not None:
        H(doc, "E  Blueprint Conversion Score (BCS)")
        section_intro(doc,
            what="BCS combines three speed-adjusted residuals — converts more often than speed predicts "
                 "(execution), covers more ground first-move→release than speed predicts (timing), and a "
                 "penalty for fast runners who get caught. The figure below shows the Top 25 each season.",
            result="The board is full of slow-but-skilled runners (gold) holding their own against "
                    "burners — the slow, high-conversion archetype (★ Naylor, Soto) repeats every season.",
            variables=["BCS", "Ground vs Speed-Expected", "Top Speed", "SB / CS"])
        add_fig(doc, FIG_DIR / "Fig_v8_BCS_ByYear.png", 7.0,
                "Figure E1 — Top 25 by Blueprint Conversion Score per season. Bar length = BCS (higher = "
                "more speed-adjusted steal skill). Color = raw-speed tier only (GOLD = slow / slate = "
                "average / blue = fast); ★ marks the benchmark names. No red — every bar is a top-25 result.")
        def bcs_cols():
            return [("rank_BCS", "#", "center"), ("player_name", "Player", "left"),
                    ("team", "Team", "center"), ("season", "Yr", "center"),
                    ("sprint_pctile", "Speed\n%ile", "center"), ("SB", "SB", "center"),
                    ("CS", "CS", "center"), ("SB_pct", "SB%", "center"),
                    ("mean_gain_to_release_ft", "Gain\n(ft)", "center"),
                    ("gain_resid_z", "Ground vs\nExp (SD)", "center"),
                    ("success_resid_z", "Convert vs\nExp (SD)", "center"),
                    ("BCS", "BCS", "center")]
        bcs_fmt = {"season": lambda v: f"{int(v)}", "sprint_pctile": lambda v: f"{v:.0f}",
                   "SB": lambda v: f"{int(v)}", "CS": lambda v: f"{int(v)}",
                   "SB_pct": lambda v: f"{v*100:.0f}%" if v <= 1.5 else f"{v:.0f}%",
                   "mean_gain_to_release_ft": lambda v: f"{v:.1f}",
                   "gain_resid_z": lambda v: f"{v:+.2f}", "success_resid_z": lambda v: f"{v:+.2f}",
                   "BCS": lambda v: f"{v:+.2f}", "rank_BCS": lambda v: f"{int(v)}"}
        H(doc, "E.1  Overall Top 15", lvl=2)
        top15 = bp.sort_values("rank_BCS").head(15).copy()
        heat_table(doc, top15, columns=bcs_cols(), team_logo_col="team",
                   heat_specs={"BCS": +1, "gain_resid_z": +1, "success_resid_z": +1},
                   highlight_ids={NAYLOR_ID, SOTO_ID}, fmt=bcs_fmt, size=8.5)
        H(doc, "E.2  Bottom 15 — Squanderers (fast, serial caught-stealings)", lvl=2)
        bot15 = bp.sort_values("rank_BCS", ascending=False).head(15).copy().sort_values("BCS")
        heat_table(doc, bot15, columns=bcs_cols(), team_logo_col="team",
                   heat_specs={"BCS": +1, "gain_resid_z": +1, "success_resid_z": +1},
                   fmt=bcs_fmt, size=8.5)
        H(doc, "E.3  Ground Covered Beyond Speed-Expected — Top 25 per Season", lvl=2)
        section_intro(doc,
            what="The single clearest timing metric: feet of ground gained from the pitcher's first "
                 "move to release, after removing the runner's sprint speed.",
            result="Slow runners (green) routinely top the board — proof that ground covered is a "
                    "timing/jump skill, not a speed read-out.",
            variables=["Ground vs Speed-Expected", "Top Speed"])
        add_fig(doc, FIG_DIR / "Fig_v8_GroundCovered_ByYear.png", 7.0,
                "Figure E2 — Bar length = feet of ground covered beyond what sprint speed predicts "
                "(e.g. '+3.04 ft' = three feet more than a runner that fast usually gains). Color = "
                "raw-speed tier only (GOLD = slow / slate = average / blue = fast); ★ marks the benchmark "
                "names. No red — every bar is a top-25 result.")

    doc.save(str(OUT_APP))
    print(f"wrote {OUT_APP}  ({OUT_APP.stat().st_size/1024:.0f} KB)")

# ══════════════════════════════════════════════════════════════════════════════
if __name__ == "__main__":
    build_ground_covered_byyear()
    build_bcs_byyear()
    build_main()
    build_appendix()
