#!/usr/bin/env python3
"""
build_v7_report.py — Comprehensive v7 report (DOCX).

One unified deliverable that absorbs:
  • v7 model sections   — Models A/B/C, AUC, GLM weights, feature importance, SSSI, xSB
  • v6 Blueprint elements — Blueprint Conversion Score, §X.3 Naylor/Soto archetype
                            profile (year-by-year), per-season Top-25 with team logos

Rankings are rendered as TABLES with an SD heat-map (green = better, grey/white =
neutral, red = worse).  Team logos are inset into the Team column of the BCS tables.

Run:  python3 build_v7_report.py   (reads existing CSVs / xlsx / figures; no network)
"""
from pathlib import Path
import pandas as pd
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT     = Path(__file__).resolve().parent
DATA_DIR = ROOT / "Data Frame"
FIG_DIR  = ROOT / "Figures"
LOGO_DIR = ROOT / "Team Logos"
REPORTS  = ROOT / "Reports"
OUT      = REPORTS / "Naylor_Model_v7_Report.docx"

NAYLOR_ID, SOTO_ID = 647304, 665742

# ── load data ────────────────────────────────────────────────────────────────
sssi = pd.read_csv(DATA_DIR / "DF_v7_SSSI.csv")
xsb  = pd.read_csv(DATA_DIR / "DF_v7_xSB_Outcome.csv")
auc  = pd.read_csv(DATA_DIR / "DF_v7_ModelB_AUC.csv")
bl   = pd.ExcelFile(DATA_DIR / "Naylor Blueprint.xlsx")
bp   = bl.parse("Blueprint Leaderboard")

# ── colour helpers ───────────────────────────────────────────────────────────
def z_to_hex(z):
    """Diverging white-centred heat-map.  z>0 → green, z<0 → red, ~0 → near-white."""
    z = max(-1.6, min(1.6, float(z)))
    if z >= 0:
        f = z / 1.6
        r = int(255 + f * (76  - 255)); g = int(255 + f * (175 - 255)); b = int(255 + f * (95  - 255))
    else:
        f = -z / 1.6
        r = int(255 + f * (214 - 255)); g = int(255 + f * (90  - 255)); b = int(255 + f * (77  - 255))
    return f"{r:02X}{g:02X}{b:02X}"

def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)

def _set_cell_text(cell, text, bold=False, size=8.5, align="left", color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER,
                   "right": WD_ALIGN_PARAGRAPH.RIGHT}[align]
    r = p.add_run(str(text)); r.bold = bold; r.font.size = Pt(size)
    if color: r.font.color.rgb = RGBColor.from_string(color)

def heat_table(doc, df, columns, heat_specs, *, team_logo_col=None,
               highlight_ids=None, fmt=None, size=8.5):
    """
    columns    : list of (df_col, header, align)
    heat_specs : {df_col: +1 | -1}  (+1 higher=greener, -1 lower=greener)
                 colour is from a z-score computed across the displayed rows.
    fmt        : {df_col: callable(value)->str}
    """
    fmt = fmt or {}
    highlight_ids = highlight_ids or set()
    # precompute per-column z over displayed rows
    zmap = {}
    for c, sign in heat_specs.items():
        v = pd.to_numeric(df[c], errors="coerce")
        sd = v.std(ddof=0)
        zmap[c] = ((v - v.mean()) / sd * sign) if sd and not np.isnan(sd) else v * 0

    tbl = doc.add_table(rows=1, cols=len(columns))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"
    for j, (_, header, align) in enumerate(columns):
        _set_cell_text(tbl.rows[0].cells[j], header, bold=True, size=size, align=align)
        shade(tbl.rows[0].cells[j], "2C3E50")
        tbl.rows[0].cells[j].paragraphs[0].runs[0].font.color.rgb = RGBColor.from_string("FFFFFF")

    for _, row in df.iterrows():
        cells = tbl.add_row().cells
        is_anchor = ("runner_id" in df.columns and row.get("runner_id") in highlight_ids)
        for j, (col, _, align) in enumerate(columns):
            if col == team_logo_col:
                _add_team_cell(cells[j], str(row[col]))
            else:
                val = row[col]
                txt = fmt[col](val) if col in fmt else (
                    f"{val:.0f}" if isinstance(val, (int, np.integer)) else
                    (f"{val:.2f}" if isinstance(val, (float, np.floating)) else str(val)))
                _set_cell_text(cells[j], txt, bold=is_anchor, size=size, align=align)
            if col in zmap:
                shade(cells[j], z_to_hex(zmap[col].loc[row.name]))
            elif is_anchor:
                shade(cells[j], "FFF3C4")   # soft amber for Naylor/Soto rows
    return tbl

def _add_team_cell(cell, team):
    cell.text = ""
    p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo = LOGO_DIR / f"{team}.png"
    if logo.exists():
        try:
            p.add_run().add_picture(str(logo), width=Inches(0.22))
            return
        except Exception:
            pass
    r = p.add_run(str(team)); r.font.size = Pt(8.5)

# ── doc-level helpers ────────────────────────────────────────────────────────
def add_fig(doc, path, width=6.6, caption=None):
    path = Path(path)
    if not path.exists():
        doc.add_paragraph(f"[missing figure: {path.name}]")
        return
    doc.add_picture(str(path), width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        c = doc.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = c.add_run(caption); r.italic = True; r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor.from_string("666666")

def H(doc, text, lvl=1):
    doc.add_heading(text, level=lvl)

def body(doc, text, size=10):
    p = doc.add_paragraph()
    r = p.add_run(text); r.font.size = Pt(size)
    return p

# ══════════════════════════════════════════════════════════════════════════════
doc = Document()
for s in ("Normal",):
    doc.styles[s].font.name = "Calibri"; doc.styles[s].font.size = Pt(10.5)

# Title
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("The Naylor Model — v7"); r.bold = True; r.font.size = Pt(26)
sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Speed-adjusted base-stealing skill · SSSI · Expected SB Outcome · Blueprint Conversion Score")
r.italic = True; r.font.size = Pt(11); r.font.color.rgb = RGBColor.from_string("666666")
doc.add_paragraph()

body(doc,
     "José Caballero led MLB in net stolen bases in 2025 running a quarter-second slower than "
     "Chandler Simpson. Josh Naylor stole 20 bases above average at 93.8% while running slower than "
     "97% of the league. Sprint speed is the most intuitive base-stealing metric — it is not the most "
     "essential one. What separates these runners is technique: lead distance, secondary-lead timing, "
     "and first-step burst. v7 measures that skill three ways — the SSSI (slow-but-skilled), the "
     "Expected SB Outcome quadrant (fast and productive), and the Blueprint Conversion Score "
     "(timing the pitcher).")

# ── §1 The Models ─────────────────────────────────────────────────────────────
H(doc, "1  The Models (A / B / C)")
mb = {r["label"]: r for _, r in auc.iterrows()} if "label" in auc.columns else {}
full = auc.loc[auc.get("label", auc.iloc[:, 0]).astype(str).str.contains("full", case=False)]
full_auc = float(full["auc"].iloc[0]) if len(full) else float("nan")
tbl = doc.add_table(rows=1, cols=4); tbl.style = "Table Grid"; tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, h in enumerate(["Model", "Unit", "AUC", "Purpose"]):
    _set_cell_text(tbl.rows[0].cells[j], h, bold=True, size=9, color="FFFFFF"); shade(tbl.rows[0].cells[j], "2C3E50")
rows = [("Model A — per-attempt GBM", "Individual attempt", "~0.59", "Strict noise-floor test"),
        ("Model B — season GBM", "Runner-season", f"{full_auc:.3f} (de-leaked)", "Headline predictor"),
        ("Model C — interpretable GLM", "Runner-season", "—", "Plain-English weight table")]
for rdat in rows:
    c = tbl.add_row().cells
    for j, v in enumerate(rdat): _set_cell_text(c[j], v, size=9)
body(doc,
     "AUC caveat (v7 de-leaking). Versions v4–v6 reported AUCs of ~0.66–0.70, but those runs carried "
     "duplicate runner-season rows — repeated Statcast split measurements for one player-season — that "
     "leaked across cross-validation folds and inflated the score. v7 averages those duplicate splits "
     "into one row per runner-season. The v7 AUC is lower but honest, not a regression. The historical "
     "bars below are kept for context only.")
add_fig(doc, FIG_DIR / "Fig_v7_AUC.png", 6.0,
        "Figure 1 — Model AUC across versions. The v7 bar is the de-leaked figure; v4–v6 are optimistic.")

# ── §2 Model C — GLM weights ──────────────────────────────────────────────────
H(doc, "2  Model C — What Actually Moves the Needle")
body(doc, "Each bar is the predicted percentage-point change in SB success when a runner improves on "
          "that feature by one tier (1 SD). Green helps; red hurts.")
add_fig(doc, FIG_DIR / "Fig_v7_GLM_PlainEnglish.png", 6.6, "Figure 2 — GLM plain-English weight table.")

# ── §3 Feature importance ─────────────────────────────────────────────────────
H(doc, "3  Feature Importance — Pre vs Post 2023")
add_fig(doc, FIG_DIR / "Fig_v7_Importance_PrePost.png", 6.2,
        "Figure 3 — How feature importance shifted after the 2023 bigger-bases rule change.")

# ── §4 SSSI ───────────────────────────────────────────────────────────────────
H(doc, "4  SSSI — Slow-Steal Skill Index")
body(doc, "A weighted composite of nine z-scored features, optimised on 80% of runners with Naylor and "
          "Soto held out entirely — so their ranking is genuinely out-of-sample. The index surfaces "
          "elite-performing slow runners. Naylor holds the top two seasons; Soto is top five.")
ss = sssi.sort_values("SSSI_v7", ascending=False).head(25).copy()
heat_table(doc, ss,
    columns=[("rank_v7", "#", "center"), ("player_name", "Player", "left"),
             ("season", "Yr", "center"), ("sprint_speed", "Sprint", "center"),
             ("jump_time", "Jump", "center"), ("sb_residual", "SB resid", "center"),
             ("accel_topspeed_premium", "Accel prem", "center"),
             ("real_sb_pct", "SB%", "center"), ("SSSI_v7", "SSSI", "center")],
    heat_specs={"sb_residual": +1, "accel_topspeed_premium": +1, "real_sb_pct": +1, "SSSI_v7": +1},
    highlight_ids={NAYLOR_ID, SOTO_ID},
    fmt={"season": lambda v: f"{int(v)}", "sprint_speed": lambda v: f"{v:.1f}",
         "real_sb_pct": lambda v: f"{v*100:.0f}%", "jump_time": lambda v: f"{v:.2f}",
         "SSSI_v7": lambda v: f"{v:+.2f}", "sb_residual": lambda v: f"{v:+.3f}",
         "accel_topspeed_premium": lambda v: f"{v:+.1f}", "rank_v7": lambda v: f"{int(v)}"})
c = doc.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
rr = c.add_run("Table 1 — SSSI v7 Top 25 (green = better on that metric, red = worse). "
               "Amber rows = Naylor / Soto."); rr.italic = True; rr.font.size = Pt(8.5)
add_fig(doc, FIG_DIR / "Fig_v7_NaylorSoto_Profile.png", 6.4,
        "Figure 4 — Naylor & Soto skill profile vs league.")

# ── §5 xSB ────────────────────────────────────────────────────────────────────
H(doc, "5  Expected SB Outcome (xSB)")
body(doc, "xSB = z(net SB above average) + z(sprint speed). A complementary lens to the SSSI: where the "
          "SSSI surfaces slow-but-skilled stealers, xSB surfaces the runners who are both fast AND "
          "productive. The companion sb_potential_gap = z(sprint) − z(net SB) splits the league into four "
          "quadrants. Positive gap = fast but under-stealing — the clearest coaching targets.")
add_fig(doc, FIG_DIR / "Fig_v7_xSB_Quadrant.png", 6.4,
        "Figure 5 — Speed-vs-production quadrant. Green = Realized Burners, blue = Untapped Wheels, "
        "red = Naylor/Soto (Crafty Technicians).")

H(doc, "5.1  Top 15 by xSB — Realized Burners (fast + productive)", lvl=2)
xt = xsb.sort_values("xsb_outcome", ascending=False).head(15).copy()
heat_table(doc, xt,
    columns=[("rank_xsb", "#", "center"), ("player_name", "Player", "left"),
             ("season", "Yr", "center"), ("sprint_speed", "Sprint", "center"),
             ("SB", "SB", "center"), ("CS", "CS", "center"), ("net_sb", "Net SB", "center"),
             ("z_sprint", "z spd", "center"), ("z_net_sb", "z stl", "center"),
             ("xsb_outcome", "xSB", "center")],
    heat_specs={"z_sprint": +1, "z_net_sb": +1, "xsb_outcome": +1, "net_sb": +1},
    fmt={"season": lambda v: f"{int(v)}", "sprint_speed": lambda v: f"{v:.1f}",
         "SB": lambda v: f"{int(v)}", "CS": lambda v: f"{int(v)}", "net_sb": lambda v: f"{int(v)}",
         "z_sprint": lambda v: f"{v:+.2f}", "z_net_sb": lambda v: f"{v:+.2f}",
         "xsb_outcome": lambda v: f"{v:+.2f}", "rank_xsb": lambda v: f"{int(v)}"})
body(doc, "Table 2 — the recognizable burners: fast and converting in volume.", 8.5)

H(doc, "5.2  Untapped Wheels — fast but under-stealing (coach them up)", lvl=2)
up = xsb[xsb["quadrant"] == "Untapped Wheels"].sort_values("sb_potential_gap", ascending=False).head(12).copy()
up.insert(0, "ut_rank", range(1, len(up) + 1))
heat_table(doc, up,
    columns=[("ut_rank", "#", "center"), ("player_name", "Player", "left"),
             ("season", "Yr", "center"), ("sprint_speed", "Sprint", "center"),
             ("SB", "SB", "center"), ("CS", "CS", "center"),
             ("z_sprint", "z spd", "center"), ("sb_potential_gap", "Potential gap", "center")],
    heat_specs={"z_sprint": +1, "sb_potential_gap": +1},
    fmt={"season": lambda v: f"{int(v)}", "sprint_speed": lambda v: f"{v:.1f}",
         "SB": lambda v: f"{int(v)}", "CS": lambda v: f"{int(v)}",
         "z_sprint": lambda v: f"{v:+.2f}", "sb_potential_gap": lambda v: f"{v:+.2f}",
         "ut_rank": lambda v: f"{int(v)}"})
body(doc, "Table 3 — biggest gaps between speed and steal production. These are the unrealized "
          "high-ceiling runners.", 8.5)

H(doc, "5.3  Naylor & Soto — Crafty Technicians", lvl=2)
ns = xsb[xsb["runner_id"].isin([NAYLOR_ID, SOTO_ID])].sort_values("xsb_outcome", ascending=False).copy()
ns.insert(0, "ns_rank", range(1, len(ns) + 1))
heat_table(doc, ns,
    columns=[("player_name", "Player", "left"), ("season", "Yr", "center"),
             ("sprint_speed", "Sprint", "center"), ("SB", "SB", "center"), ("CS", "CS", "center"),
             ("z_sprint", "z spd", "center"), ("z_net_sb", "z stl", "center"),
             ("sb_potential_gap", "Potential gap", "center"), ("quadrant", "Quadrant", "left")],
    heat_specs={"z_net_sb": +1, "sb_potential_gap": -1},
    highlight_ids={NAYLOR_ID, SOTO_ID},
    fmt={"season": lambda v: f"{int(v)}", "sprint_speed": lambda v: f"{v:.1f}",
         "SB": lambda v: f"{int(v)}", "CS": lambda v: f"{int(v)}",
         "z_sprint": lambda v: f"{v:+.2f}", "z_net_sb": lambda v: f"{v:+.2f}",
         "sb_potential_gap": lambda v: f"{v:+.2f}"})
body(doc, "Table 4 — negative potential gap = they out-steal their speed. The archetype this whole "
          "model is built around.", 8.5)

# ── §6 Blueprint Conversion Score ─────────────────────────────────────────────
H(doc, "6  Blueprint Conversion Score (BCS)")
body(doc, "BCS combines three speed-adjusted residuals: success_resid_z (converts more often than "
          "sprint speed predicts — execution), gain_resid_z (covers more ground first-move→release "
          "than speed predicts — timing), and squander_z (penalty for fast runners who get caught). "
          "It scores the Naylor archetype directly: slow sprint, big jump, elite conversion.")

def bcs_cols(extra_team=True):
    cols = [("rank_BCS", "#", "center"), ("player_name", "Player", "left")]
    if extra_team: cols.append(("team", "Team", "center"))
    cols += [("season", "Yr", "center"), ("sprint_pctile", "Spd %ile", "center"),
             ("SB", "SB", "center"), ("CS", "CS", "center"), ("SB_pct", "SB%", "center"),
             ("mean_gain_to_release_ft", "Gain ft", "center"),
             ("gain_resid_z", "Gn z", "center"), ("success_resid_z", "Su z", "center"),
             ("BCS", "BCS", "center")]
    return cols

bcs_fmt = {"season": lambda v: f"{int(v)}", "sprint_pctile": lambda v: f"{v:.0f}",
           "SB": lambda v: f"{int(v)}", "CS": lambda v: f"{int(v)}",
           "SB_pct": lambda v: f"{v*100:.0f}%" if v <= 1.5 else f"{v:.0f}%",
           "mean_gain_to_release_ft": lambda v: f"{v:.1f}",
           "gain_resid_z": lambda v: f"{v:+.2f}", "success_resid_z": lambda v: f"{v:+.2f}",
           "BCS": lambda v: f"{v:+.2f}", "rank_BCS": lambda v: f"{int(v)}"}

H(doc, "6.1  Overall Top 15 — The Blueprint", lvl=2)
top15 = bp.sort_values("rank_BCS").head(15).copy()
heat_table(doc, top15, columns=bcs_cols(), team_logo_col="team",
           heat_specs={"BCS": +1, "gain_resid_z": +1, "success_resid_z": +1},
           highlight_ids={NAYLOR_ID, SOTO_ID}, fmt=bcs_fmt)
body(doc, "Table 5 — Overall Top 15 BCS (2023–2026). Logos mark each runner's team.", 8.5)
add_fig(doc, FIG_DIR / "Fig_NaylorBlueprint_TopN.png", 6.4,
        "Figure 6 — Top BCS runner-seasons. Red = Naylor, green = Soto.")

H(doc, "6.2  Naylor & Soto — The Archetype Profile (year-by-year)", lvl=2)
ar = bp[bp["runner_id"].isin([NAYLOR_ID, SOTO_ID])].sort_values(["player_name", "season"]).copy()
heat_table(doc, ar, columns=bcs_cols(), team_logo_col="team",
           heat_specs={"BCS": +1, "gain_resid_z": +1, "success_resid_z": +1},
           highlight_ids={NAYLOR_ID, SOTO_ID}, fmt=bcs_fmt)
body(doc, "Table 6 — Naylor & Soto, all tracked seasons.", 8.5)
# dynamic year-by-year prose
lines = []
for _, r in ar.iterrows():
    last = r["player_name"].split()[-1]
    lines.append(f"{last} {int(r['season'])}: {r['sprint_speed_ftps']:.1f} ft/s "
                 f"({r['sprint_pctile']:.0f}th pct), {r['mean_gain_to_release_ft']:.1f} ft gain, "
                 f"{int(r['SB'])}/{int(r['SB'])+int(r['CS'])} "
                 f"= {r['SB_pct']*100 if r['SB_pct']<=1.5 else r['SB_pct']:.0f}% SB, BCS {r['BCS']:+.2f}.")
for ln in lines:
    p = doc.add_paragraph(style="List Bullet"); rr = p.add_run(ln); rr.font.size = Pt(9.5)
body(doc, "The Naylor archetype — slow sprint, big jump, high conversion — appears in every season and "
          "is not a single-year outlier. Naylor 2026 is already tracking above his 2025 pace.")
add_fig(doc, FIG_DIR / "Fig_NaylorBlueprint_Scatter.png", 6.0,
        "Figure 7 — BCS scatter: sprint speed vs mean gain, coloured by score. Naylor/Soto annotated.")

H(doc, "6.3  Top 25 Per Season — Ground Covered", lvl=2)
add_fig(doc, FIG_DIR / "Fig_GroundCovered_TopN_ByYear.png", 6.8,
        "Figure 8 — Top 25 per season by speed-adjusted ground covered, with team logos.")

H(doc, "6.4  Bottom 15 — The Anti-Naylor (Squanderers)", lvl=2)
bot15 = bp.sort_values("rank_BCS", ascending=False).head(15).copy()
bot15 = bot15.sort_values("BCS")  # worst first
heat_table(doc, bot15, columns=bcs_cols(), team_logo_col="team",
           heat_specs={"BCS": +1, "gain_resid_z": +1, "success_resid_z": +1}, fmt=bcs_fmt)
body(doc, "Table 7 — fast runners who get caught repeatedly. Bobby Witt Jr. is bottom-five every "
          "season in the dataset.", 8.5)
add_fig(doc, FIG_DIR / "Fig_NaylorBlueprint_BottomN.png", 6.4,
        "Figure 9 — Bottom BCS runner-seasons: elite speed, serial caught-stealings.")

H(doc, "6.5  Per-Season BCS Top 25 — with Team Logos", lvl=2)
add_fig(doc, FIG_DIR / "Fig_BCS_Top25_ByYear.png", 6.9,
        "Figure 10 — Top 25 BCS per season (2023–2026) with team logos. Red = Naylor, green = Soto.")

# ── §7 Conclusions ────────────────────────────────────────────────────────────
H(doc, "7  Conclusions")
for ln in [
    "Ground covered (first move → release) is near speed-independent — sprint speed explains only a "
    "small fraction of the variance (r = −0.36).",
    "The gain metric meaningfully predicts steal success (r ≈ +0.25–0.29 with SB%). The signal is real.",
    "The Naylor archetype — slow sprint, big jump, high conversion — is stable across every season.",
    "The squander archetype — fast sprint, serial caught-stealings — is equally stable.",
    "Practical implication: timing and jump training can shift scores at both ends of the speed "
    "spectrum. Sprint speed is structural; the jump window is behavioral.",
]:
    p = doc.add_paragraph(style="List Bullet"); rr = p.add_run(ln); rr.font.size = Pt(10)

doc.save(str(OUT))
print(f"wrote {OUT}  ({OUT.stat().st_size/1024:.0f} KB)")
